from __future__ import annotations

import shutil
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = Path("/Users/gireesha/Downloads/agent_portfolio_summary_api_spec.md")
SOURCE_YAML = Path("/Users/gireesha/Downloads/agent_portfolio_summary_api_spec.yaml")

OUT_DOCX = ROOT / "docs" / "Agent_Portfolio_Summary_API_Spec_World_Class.docx"
OUT_FULL_DOCX = ROOT / "docs" / "Agent_Portfolio_Summary_API_Spec_Full_With_Examples.docx"
OUT_MD = ROOT / "docs" / "Agent_Portfolio_Summary_API_Spec_World_Class.md"
OUT_YAML = ROOT / "docs" / "Agent_Portfolio_Summary_API_OpenAPI_Enhanced.yaml"
OUT_SOURCE_MD = ROOT / "docs" / "Agent_Portfolio_Summary_API_Source_Reference.md"

NAVY = "1F4E79"
TEAL = "0F766E"
SLATE = "334155"
LIGHT_BLUE = "EAF3F8"
LIGHT_TEAL = "E7F6F3"
LIGHT_GRAY = "F6F8FA"
WHITE = "FFFFFF"

DOCUMENT_META = [
    ("Document", "Agent Portfolio Summary API Specification"),
    ("Version", "2.0 - enhanced from attached Markdown and OpenAPI YAML"),
    ("Prepared on", "19 May 2026"),
    ("Source Markdown", str(SOURCE_MD)),
    ("Source OpenAPI", str(SOURCE_YAML)),
    ("API base path", "/sbos-ibpm/v1"),
    ("Primary resource", "Agent / Business Correspondent portfolio summary"),
    ("Primary key", "agentCode"),
]

ENDPOINTS = [
    (
        "GET",
        "/agents/{agentCode}/portfolio-summary",
        "getAgentPortfolioSummary",
        "Return complete dashboard-ready portfolio summary for an agent across all requested windows and sections.",
        "agentCode path parameter; windowCodes, includeTrends, includeBreakdowns, includeNarratives query parameters.",
        "AgentPortfolioSummaryResponse",
    ),
    (
        "GET",
        "/agents/{agentCode}/portfolio-summary/sections/{sectionCode}",
        "getAgentPortfolioSection",
        "Return one selected section, such as Transaction Summary, Commission Summary, SPM, TCM, ARM, AAM, ADRM, or EAPRM.",
        "agentCode and sectionCode path parameters; windowCodes, includeTrends, includeBreakdowns query parameters.",
        "PortfolioSectionResponse",
    ),
    (
        "GET",
        "/agents/{agentCode}/portfolio-trends",
        "getAgentPortfolioTrends",
        "Return chart-ready trend series for selected sections using DAY, MONTH, or QUARTER grain.",
        "agentCode path parameter; grain and sectionCodes query parameters; optional fromDate and toDate.",
        "AgentPortfolioTrendResponse",
    ),
    (
        "POST",
        "/agents/{agentCode}/portfolio-summary/query",
        "queryAgentPortfolioSummary",
        "Advanced query for dashboards requiring section filtering, hierarchy filters, channel filters, transaction/service filters, and custom date range.",
        "agentCode path parameter; PortfolioSummaryQueryRequest JSON body.",
        "AgentPortfolioSummaryResponse",
    ),
    (
        "GET",
        "/reference/portfolio-codes",
        "getPortfolioReferenceCodes",
        "Return supported window, grain, and section codes for client configuration.",
        "No request body.",
        "PortfolioReferenceResponse",
    ),
]

QUERY_PARAMS = [
    ("windowCodes", "query string", "No", "LAST_7_DAYS,LAST_3_MONTHS,LAST_4_QUARTERS", "Comma-separated window codes to include."),
    ("includeTrends", "boolean", "No", "true", "Include trendSummary.periods[] for charts and period drill-down."),
    ("includeBreakdowns", "boolean", "No", "true", "Include breakdowns by channel, transaction type, service type, and applicable product type."),
    ("includeNarratives", "boolean", "No", "true", "Include simple business explanation messages for agent-facing screens."),
    ("grain", "enum", "Trend endpoint only", "DAY", "DAY, MONTH, or QUARTER."),
    ("sectionCodes", "query string", "Trend endpoint only", "TRANSACTION_SUMMARY,COMMISSION_SUMMARY", "Comma-separated section codes."),
    ("fromDate", "date", "Custom trend only", "2026-05-13", "Optional custom period start date."),
    ("toDate", "date", "Custom trend only", "2026-05-19", "Optional custom period end date."),
]

WINDOWS = [
    ("LAST_7_DAYS", "Seven day date-wise trend", "DAY", "7", "Date-wise charts and daily agent action tracking."),
    ("LAST_3_MONTHS", "Three month trend", "MONTH", "3", "Month-wise comparison and recent performance trajectory."),
    ("LAST_4_QUARTERS", "Four quarter trend", "QUARTER", "4", "Quarter-wise business and FY comparison."),
    ("CUSTOM", "Custom analytics range", "DAY / MONTH / QUARTER", "Request dependent", "Filtered analysis using explicit fromDate and toDate."),
    ("LAST_30_DAYS", "Optional compatibility code in source YAML", "DAY", "30", "Can be retained in reference data if existing clients already use it."),
]

SECTIONS = [
    (
        "TRANSACTION_SUMMARY",
        "transactionSummary",
        "Transaction Summary",
        "Financial/non-financial transaction activity, enrollments, lead services, ease services, cash movement, ratios, and channel/type/service breakdowns.",
        "Trend chart, cash in/out split, channel and service filters.",
    ),
    (
        "COMMISSION_SUMMARY",
        "commissionSummary",
        "Commission Summary",
        "SBOS Autopay commission earned, eligible transactions, eligible value, rank, rank movement, and average commission.",
        "Commission card, rank movement, monthly commission trend.",
    ),
    (
        "SYSTEM_PERFORMANCE",
        "systemPerformance",
        "System Performance",
        "Success/failure ratio, business declines, technical declines, top BD/TD response code, and health band.",
        "Success gauge, error trend, top decline reason panel.",
    ),
    (
        "TARGET_CAMPAIGN_MANAGEMENT",
        "targetCampaignManagement",
        "Target & Campaign Management",
        "Assigned targets, achievement percentage, pending transactions, remaining days, target status, and required daily pace.",
        "Progress bar, required-per-day counter, target drill-down.",
    ),
    (
        "ANOMALY_RISK_MANAGEMENT",
        "anomalyRiskManagement",
        "Anomaly Risk Management",
        "Suspicious transaction ratio, suspicious customers, suspicious amount, anomaly cases, top anomaly, and risk band.",
        "Risk badge, anomaly panel, suspicious amount trend.",
    ),
    (
        "AGENT_AUDIT_MANAGEMENT",
        "agentAuditManagement",
        "Agent Audit Management",
        "Audits initiated, marks scored, ongoing audit count, recent audits, audit source, audit type, and last audited date.",
        "Audit timeline, open audit badge, audit quality band.",
    ),
    (
        "ANALYTICS_REPORTING_MANAGEMENT",
        "analyticsReportingManagement",
        "Analytics Reporting Management",
        "Descriptive, diagnostic, predictive, and prescriptive recommendations with severity and lifecycle status.",
        "Recommendation feed grouped by module and priority.",
    ),
    (
        "ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT",
        "enterpriseAgentPerformanceRiskManagement",
        "Enterprise Agent Performance Risk Management",
        "EAPRM iScore, band, operational, behavioural, compliance, financial health, customer engagement, and attendance scores.",
        "iScore hero card, component score bars, score trend.",
    ),
]

RESPONSE_ENVELOPE = [
    ("requestId", "string", "Yes", "Unique request/correlation id."),
    ("schemaVersion", "string", "Yes", "Response schema version. Recommended value: 2.0."),
    ("referenceDataVersion", "string", "Yes", "Taxonomy version for windows, sections, channels, transaction types, and service types."),
    ("generatedAt", "datetime", "Yes", "API response generation timestamp."),
    ("agent", "object", "Yes", "Agent identity, branch, region, status, and hierarchy level."),
    ("portfolioStatus", "object", "Yes", "Overall band, refresh status, data freshness, and warning."),
    ("timeWindows[]", "array", "Yes", "One object per requested window. Each object contains sections."),
    ("timeWindows[].sections", "object", "Yes", "Map containing selected portfolio sections."),
    ("recommendations[]", "array", "No", "Cross-module recommendation list."),
]

SECTION_PATTERN = [
    ("overall", "object", "Dashboard-card summary values for the selected window."),
    ("trendSummary.grain", "DAY / MONTH / QUARTER", "The period grain used by trendSummary.periods[]."),
    ("trendSummary.periods[]", "array", "Chart and drill-down rows. Every row must include periodCode, periodLabel, grain, fromDate, and toDate."),
    ("breakdowns", "object", "Optional channel/type/service/product breakdowns when includeBreakdowns=true."),
    ("narrative", "string", "Optional simple explanation when includeNarratives=true."),
]

FIELD_CATALOG = [
    ("Transaction", "overall.financialTxnCount", "integer", "SUM", "Financial transaction count."),
    ("Transaction", "overall.nonFinancialTxnCount", "integer", "SUM", "Non-financial transaction count."),
    ("Transaction", "overall.enrollmentCount", "integer", "SUM", "Enrollment count for PMJJBY, PMSBY, APY, and future types."),
    ("Transaction", "overall.leadServiceCount", "integer", "SUM", "Lead service count for Loan, RD, FD, and future services."),
    ("Transaction", "overall.easeServiceCount", "integer", "SUM", "Ease service count such as debit card hotlisting and account enquiry."),
    ("Transaction", "overall.totalTxnCount", "integer", "SUM", "Total transaction count."),
    ("Transaction", "overall.amountInvolved", "Money", "SUM", "Gross amount involved in selected transactions."),
    ("Transaction", "overall.financialTxnRatio", "number", "Weighted ratio", "financialTxnCount / totalTxnCount * 100."),
    ("Transaction", "overall.averageFinancialTxnRatio", "number", "Weighted average", "Comparison average for selected context."),
    ("Transaction", "overall.cashInflow", "Money", "SUM", "Cash deposit / inflow value."),
    ("Transaction", "overall.cashOutflow", "Money", "SUM", "Cash withdrawal / outflow value."),
    ("Transaction", "breakdowns.byChannel[]", "array", "Group by", "ONUS, OFFUS, and future configured channels."),
    ("Transaction", "breakdowns.byTransactionType[]", "array", "Group by", "AEPS, RUPAY, SHG, TPD, IMPS, and future types."),
    ("Transaction", "breakdowns.byServiceType[]", "array", "Group by", "CD, CW, FT, and future service types."),
    ("Commission", "overall.totalCommissionEarned", "Money", "SUM", "Commission earned/accrued."),
    ("Commission", "overall.totalEligibleTxnCount", "integer", "SUM", "Commission-eligible transaction count."),
    ("Commission", "overall.eligibleTxnAmount", "Money", "SUM", "Eligible transaction value."),
    ("Commission", "overall.commissionRank", "RankMovement", "Latest + movement", "Rank inside peer group."),
    ("Commission", "overall.totalTxnCount", "integer", "SUM", "Transactions considered by commission rules."),
    ("Commission", "overall.averageCommissionEarned", "Money", "Weighted average", "Commission per eligible transaction."),
    ("System Performance", "overall.failureRatio", "number", "Weighted ratio", "Failed transactions / total transactions * 100."),
    ("System Performance", "overall.successRatio", "number", "Weighted ratio", "Successful transactions / total transactions * 100."),
    ("System Performance", "overall.businessDeclineCount", "integer", "SUM", "Business decline count."),
    ("System Performance", "overall.technicalDeclineCount", "integer", "SUM", "Technical decline count."),
    ("System Performance", "overall.topBusinessDeclineResponseCode", "object", "Mode", "Most frequent BD response code."),
    ("System Performance", "overall.topTechnicalDeclineResponseCode", "object", "Mode", "Most frequent TD response code."),
    ("TCM", "overall.totalTargetsAssigned", "integer", "SUM", "Assigned target count."),
    ("TCM", "overall.totalTargetAchievement", "number", "Weighted ratio", "Overall achievement percentage."),
    ("TCM", "targets[].achievementPercent", "number", "Per target ratio", "Target-wise achievement percentage."),
    ("TCM", "targets[].txnPerDayRequired", "integer", "Derived", "pendingTxnCount / remainingDays."),
    ("TCM", "targets[].pendingTxnCount", "integer", "Derived", "Remaining transaction count to meet target."),
    ("TCM", "targets[].remainingDays", "integer", "Date diff", "Days remaining in campaign."),
    ("ARM", "overall.suspiciousTxnRatio", "number", "Weighted ratio", "Suspicious transactions / total transactions * 100."),
    ("ARM", "overall.totalSuspiciousTxnCount", "integer", "SUM", "Suspicious transaction count."),
    ("ARM", "overall.totalAnomalyCaseCount", "integer", "SUM / distinct", "Anomaly cases linked to agent."),
    ("ARM", "overall.suspiciousCustomerCount", "integer", "Distinct count", "Distinct suspicious customers."),
    ("ARM", "overall.suspiciousAmount", "Money", "SUM", "Amount involved in suspicious transactions."),
    ("ARM", "overall.topAnomalyCase", "object", "Mode by severity/count", "Dominant anomaly case."),
    ("AAM", "overall.totalAuditsInitiated", "integer", "SUM", "Audit count initiated."),
    ("AAM", "overall.auditedMarksScored", "number", "Latest / weighted average", "Marks in completed audits."),
    ("AAM", "overall.ongoingAuditCount", "integer", "Current count", "Open audit count."),
    ("AAM", "overall.lastAuditedOn", "date", "MAX", "Latest audit date."),
    ("AAM", "recentAudits[]", "array", "Latest N", "Recent audit details."),
    ("ADRM", "overall.totalRecommendations", "integer", "SUM", "Recommendation count."),
    ("ADRM", "recommendations[].type", "enum", "Classified", "DESCRIPTIVE, DIAGNOSTIC, PREDICTIVE, PRESCRIPTIVE."),
    ("ADRM", "recommendations[].severity", "enum", "Classified", "Critical, Major, Minor, Info."),
    ("EAPRM", "overall.iScore", "integer", "Latest", "Enterprise iScore, 0-1000."),
    ("EAPRM", "overall.band", "string", "Band rule", "Critical, High Risk, Medium Risk, Stable, Good, Very Good."),
    ("EAPRM", "overall.operationalEfficiencyScore", "number", "Latest", "Operational score."),
    ("EAPRM", "overall.behaviouralEfficiencyScore", "number", "Latest", "Behavioral score."),
    ("EAPRM", "overall.complianceScore", "number", "Latest", "Compliance score."),
    ("EAPRM", "overall.financialHealthScore", "number", "Latest", "Financial health score."),
    ("EAPRM", "overall.totalCustomersEngaged", "integer", "Distinct count", "Customers engaged."),
    ("EAPRM", "overall.attendancePercentage", "number", "Percent", "Attendance percentage."),
]

BAND_RULES = [
    ("TCM achievement", "0%", "Not Started"),
    ("TCM achievement", "1-24%", "Critical"),
    ("TCM achievement", "25-50%", "Low"),
    ("TCM achievement", "51-75%", "Average"),
    ("TCM achievement", "76-99%", "High"),
    ("TCM achievement", "100%+", "Achieved"),
    ("ARM risk", "0%", "Genuine"),
    ("ARM risk", "1-25%", "Very Low Risk"),
    ("ARM risk", "26-50%", "Medium Risk"),
    ("ARM risk", "51-75%", "High Risk"),
    ("ARM risk", "76-99%", "Very High Risk"),
    ("ARM risk", "100%", "Fraud"),
    ("SPM error rate", "0%", "Perfect"),
    ("SPM error rate", ">0-1%", "Very Low"),
    ("SPM error rate", ">1-3%", "Low"),
    ("SPM error rate", ">3-5%", "Moderate"),
    ("SPM error rate", ">5-10%", "High"),
    ("SPM error rate", ">10%", "Critical"),
    ("AAM score", "100%", "Perfect"),
    ("AAM score", "75-99%", "High"),
    ("AAM score", "50-74%", "Moderate"),
    ("AAM score", "25-49%", "Low"),
    ("AAM score", "1-24%", "Very Low"),
    ("AAM score", "0%", "Critical"),
    ("EAPRM iScore", "<600", "Critical"),
    ("EAPRM iScore", "601-700", "High Risk"),
    ("EAPRM iScore", "701-800", "Medium Risk"),
    ("EAPRM iScore", "801-900", "Stable"),
    ("EAPRM iScore", "901-950", "Good"),
    ("EAPRM iScore", "951-1000", "Very Good"),
]

ERRORS = [
    ("400", "INVALID_WINDOW_CODE", "Unsupported window code supplied.", "Show validation message and refresh reference codes."),
    ("400", "INVALID_SECTION_CODE", "Unsupported section code supplied.", "Use /reference/portfolio-codes."),
    ("401", "UNAUTHORIZED", "Missing or invalid token.", "Re-authenticate."),
    ("403", "FORBIDDEN_AGENT_SCOPE", "Caller cannot view requested agent.", "Show access denied."),
    ("404", "AGENT_NOT_FOUND", "Agent code not found.", "Ask user to verify agent code."),
    ("409", "REFRESH_IN_PROGRESS", "Daily snapshot is running.", "Retry or show last successful snapshot."),
    ("422", "INVALID_FILTER", "Filter value is not configured.", "Refresh reference data."),
    ("429", "RATE_LIMITED", "Too many requests.", "Retry after configured wait."),
    ("500", "PORTFOLIO_SUMMARY_FAILED", "Unexpected service or aggregation failure.", "Log requestId and show friendly error."),
]

ADVANCED_QUERY_EXAMPLE = """{
  "windowCodes": ["LAST_7_DAYS", "LAST_3_MONTHS", "LAST_4_QUARTERS"],
  "sectionCodes": [
    "TRANSACTION_SUMMARY",
    "COMMISSION_SUMMARY",
    "SYSTEM_PERFORMANCE",
    "TARGET_CAMPAIGN_MANAGEMENT",
    "ANOMALY_RISK_MANAGEMENT",
    "AGENT_AUDIT_MANAGEMENT",
    "ANALYTICS_REPORTING_MANAGEMENT",
    "ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT"
  ],
  "includeTrends": true,
  "includeBreakdowns": true,
  "includeNarratives": true,
  "filters": {
    "channelCodes": ["ONUS", "OFFUS"],
    "transactionTypeCodes": ["AEPS", "RUPAY", "SHG", "TPD", "IMPS"],
    "serviceTypeCodes": ["CD", "CW", "FT"],
    "enrollmentTypeCodes": ["PMJJBY", "PMSBY", "APY"],
    "leadServiceTypeCodes": ["Loan", "RD", "FD"],
    "easeServiceTypeCodes": ["Debit Card Hotlisting", "Account Enquiry"]
  }
}"""

RESPONSE_SKELETON = """{
  "requestId": "req_01HY9Z6N0W5D4Z7D6T7E8K9P2A",
  "schemaVersion": "2.0",
  "referenceDataVersion": "svc-taxonomy-2026-05",
  "generatedAt": "2026-05-19T18:45:30+05:30",
  "agent": { "agentCode": "BC00012345", "status": "ACTIVE" },
  "portfolioStatus": {
    "overallBand": "Stable",
    "dailyRefreshStatus": "COMPLETED",
    "dataFreshness": { "latestTxnDate": "2026-05-19" }
  },
  "timeWindows": [
    {
      "code": "LAST_7_DAYS",
      "grain": "DAY",
      "sections": {
        "transactionSummary": {
          "overall": {},
          "trendSummary": { "grain": "DAY", "periods": [] },
          "breakdowns": {},
          "narrative": "Business explanation"
        }
      }
    }
  ],
  "recommendations": []
}"""


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 8.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple | list],
    widths: list[float] | None = None,
    font_size: float = 8.2,
    header_fill: str = NAVY,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, header_fill)
        set_cell_text(cell, header, bold=True, color=WHITE, size=font_size)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            set_cell_text(cells[col_index], value, size=font_size)
            if row_index % 2 == 1:
                shade_cell(cells[col_index], "F8FAFC")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)


def set_margins(section, top=0.45, bottom=0.45, left=0.52, right=0.52) -> None:
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section, top=0.42, bottom=0.42, left=0.45, right=0.45)


def set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    if section.page_width > section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Agent Portfolio Summary API Specification | v2.0"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(7.5)
        header.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
        footer = section.footer.paragraphs[0]
        footer.text = "Source-aligned enterprise API draft | 19 May 2026"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(7.5)
        footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    for name, size, color in [
        ("Heading 1", 15, NAVY),
        ("Heading 2", 12, TEAL),
        ("Heading 3", 10, SLATE),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(5 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(2)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.02
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def wrap_code(code: str, width: int = 96) -> list[str]:
    lines: list[str] = []
    for raw in code.strip("\n").splitlines():
        line = raw.rstrip()
        if len(line) <= width:
            lines.append(line)
            continue
        indent = line[: len(line) - len(line.lstrip(" "))]
        lines.extend(
            textwrap.wrap(
                line,
                width=width,
                subsequent_indent=indent + "  ",
                break_long_words=False,
                break_on_hyphens=False,
            )
        )
    return lines


def add_code(doc: Document, code: str) -> None:
    for line in wrap_code(code):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.86
        p.paragraph_format.left_indent = Inches(0.06)
        shade_paragraph(p, LIGHT_GRAY)
        run = p.add_run(line or " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(6.3)
        run.font.color.rgb = RGBColor.from_string("111827")


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    header_line = lines[start].strip()
    separator_line = lines[start + 1].strip()
    if not (header_line.startswith("|") and header_line.endswith("|")):
        return None
    separator_cells = [cell.strip() for cell in separator_line.strip("|").split("|")]
    if not separator_cells or not all(set(cell.replace(":", "").strip()) <= {"-"} and "-" in cell for cell in separator_cells):
        return None

    headers = [cell.strip().strip("`") for cell in header_line.strip("|").split("|")]
    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines):
        line = lines[index].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        row = [cell.strip().strip("`") for cell in line.strip("|").split("|")]
        if len(row) < len(headers):
            row.extend([""] * (len(headers) - len(row)))
        if len(row) > len(headers):
            row = row[: len(headers) - 1] + [" | ".join(row[len(headers) - 1 :])]
        rows.append(row)
        index += 1
    return headers, rows, index


def render_markdown_source(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    code_label = ""

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                if code_label:
                    add_para(doc, f"Example: {code_label}")
                add_code(doc, "\n".join(code_lines))
                code_lines = []
                code_label = ""
                in_code = False
            else:
                in_code = True
                code_label = stripped.strip("`").strip()
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not stripped or stripped == "---":
            index += 1
            continue

        table = parse_markdown_table(lines, index)
        if table:
            headers, rows, next_index = table
            add_table(doc, headers, rows, font_size=6.8, header_fill=TEAL)
            index = next_index
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            text = stripped.lstrip("#").strip()
            add_heading(doc, text, level=level)
        elif stripped.startswith("- "):
            add_bullets(doc, [stripped[2:].strip()])
        elif stripped[0:3].isdigit() and ". " in stripped[:5]:
            add_para(doc, stripped)
        else:
            add_para(doc, stripped)
        index += 1

    if in_code and code_lines:
        add_code(doc, "\n".join(code_lines))


def add_note(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    r.font.size = Pt(9)
    p.add_run(body)


def build_markdown() -> str:
    lines: list[str] = [
        "# Agent Portfolio Summary API Specification",
        "",
        "Version: 2.0 - enhanced from attached Markdown and OpenAPI YAML",
        "",
        "## Executive Summary",
        "This specification defines a source-aligned, enterprise-grade Agent/BC Portfolio Summary API. It preserves the attached API shape and strengthens it with clear endpoint contracts, response envelope rules, section field catalog, trend standards, security, error handling, observability, and frontend usage guidance.",
        "",
        "## Endpoints",
        "| Method | Path | Operation | Purpose |",
        "|---|---|---|---|",
    ]
    for method, path, op, purpose, _, _ in ENDPOINTS:
        lines.append(f"| `{method}` | `{path}` | `{op}` | {purpose} |")
    lines.extend(["", "## Time Windows", "| Code | Grain | Periods | Purpose |", "|---|---|---:|---|"])
    for code, purpose, grain, periods, usage in WINDOWS:
        lines.append(f"| `{code}` | `{grain}` | {periods} | {usage} |")
    lines.extend(["", "## Section Pattern", "| Field | Type | Rule |", "|---|---|---|"])
    for field, type_, rule in SECTION_PATTERN:
        lines.append(f"| `{field}` | {type_} | {rule} |")
    lines.extend(["", "## Sections", "| Section Code | JSON Field | Purpose |", "|---|---|---|"])
    for code, field, title, purpose, _ in SECTIONS:
        lines.append(f"| `{code}` | `{field}` | **{title}.** {purpose} |")
    lines.extend(["", "## Field Catalog", "| Module | Field | Type | Aggregation | Definition |", "|---|---|---|---|---|"])
    for module, field, type_, agg, definition in FIELD_CATALOG:
        lines.append(f"| {module} | `{field}` | {type_} | {agg} | {definition} |")
    lines.extend(["", "## Advanced Query Example", "```json", ADVANCED_QUERY_EXAMPLE, "```"])
    if SOURCE_MD.exists():
        lines.extend([
            "",
            "## Full API Specification and Examples From Attached Markdown",
            "",
            SOURCE_MD.read_text(encoding="utf-8"),
        ])
    if SOURCE_YAML.exists():
        lines.extend([
            "",
            "## Full OpenAPI YAML Contract",
            "",
            "```yaml",
            enhanced_yaml(),
            "```",
        ])
    return "\n".join(lines) + "\n"


def enhanced_yaml() -> str:
    source = SOURCE_YAML.read_text(encoding="utf-8") if SOURCE_YAML.exists() else ""
    enhancement = """

x-blueprnt-enhancements:
  documentationVersion: "2.0"
  generatedDocument: "docs/Agent_Portfolio_Summary_API_Spec_World_Class.docx"
  sourceAligned: true
  designRules:
    - "Every section uses overall, trendSummary.periods, breakdowns, narrative."
    - "Trend rows must include periodCode, periodLabel, grain, fromDate, toDate."
    - "Dashboard cards read from overall; charts read from trendSummary.periods."
    - "Reference data owns channel, transaction, service, enrollment, lead, and ease-service taxonomies."
    - "The API should return partial-data warnings instead of hiding stale modules."
  recommendedHeaders:
    - Authorization
    - X-Correlation-Id
    - X-Client-App
    - Accept-Language
"""
    return source.rstrip() + enhancement


def build_docx() -> None:
    doc = Document()
    set_margins(doc.sections[0])
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Agent Portfolio Summary\nAPI Specification")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    r = subtitle.add_run("Source-aligned enterprise contract for summary, sections, trends, advanced query, and reference APIs")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    add_table(doc, ["Document Field", "Value"], DOCUMENT_META, widths=[1.7, 5.7], font_size=8.4, header_fill=TEAL)

    add_heading(doc, "1. Executive Summary")
    add_para(
        doc,
        "This document enhances the attached Agent Portfolio Summary API into a compact, implementation-ready enterprise specification. It keeps the source API structure: complete summary, one-section retrieval, trend retrieval, advanced query, and reference-code retrieval."
    )
    add_para(
        doc,
        "Every portfolio section follows the same response pattern: overall for dashboard cards, trendSummary.periods[] for charts and drill-down, breakdowns for channel/type/service analysis, and narrative for agent-friendly explanation."
    )
    add_note(
        doc,
        "Document layout decision",
        "Long raw JSON/YAML is not placed in oversized table cells. API examples are short, and the complete OpenAPI YAML is exported as a separate file to prevent hidden text and unnecessary blank pages.",
        fill=LIGHT_TEAL,
    )

    add_heading(doc, "2. API Surface")
    add_table(doc, ["Method", "Path", "Operation", "Purpose", "Input", "Response"], ENDPOINTS, widths=[0.55, 2.45, 1.45, 2.45, 2.0, 1.45], font_size=6.7)

    add_heading(doc, "3. Request Parameters")
    add_table(doc, ["Parameter", "Type", "Required", "Example", "Description"], QUERY_PARAMS, widths=[1.25, 1.1, 1.1, 2.0, 3.8], font_size=7.4)

    add_heading(doc, "4. Time Window and Grain Standard")
    add_table(doc, ["Window Code", "Purpose", "Grain", "Expected Periods", "Frontend Usage"], WINDOWS, widths=[1.35, 1.85, 1.35, 1.2, 2.65], font_size=7.5)
    add_bullets(doc, [
        "LAST_7_DAYS returns exactly seven DAY rows.",
        "LAST_3_MONTHS returns exactly three MONTH rows.",
        "LAST_4_QUARTERS returns exactly four QUARTER rows.",
        "CUSTOM requires explicit fromDate and toDate and must echo the resolved period rows.",
    ])

    add_heading(doc, "5. Standard Response Envelope")
    add_table(doc, ["Field", "Type", "Required", "Description"], RESPONSE_ENVELOPE, widths=[2.2, 1.15, 0.8, 3.7], font_size=7.5)

    add_heading(doc, "6. Standard Section Pattern")
    add_table(doc, ["Field", "Type", "Contract Rule"], SECTION_PATTERN, widths=[2.1, 1.7, 4.0], font_size=7.8, header_fill=TEAL)

    add_heading(doc, "7. Portfolio Sections")
    add_table(doc, ["Section Code", "JSON Field", "Section", "Business Scope", "Frontend Representation"], SECTIONS, widths=[1.75, 1.65, 1.45, 3.0, 2.15], font_size=6.7)

    add_heading(doc, "8. Advanced Query Example")
    add_code(doc, ADVANCED_QUERY_EXAMPLE)

    add_heading(doc, "9. Response Skeleton")
    add_code(doc, RESPONSE_SKELETON)

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(landscape)
    add_header_footer(doc)
    add_heading(doc, "10. Complete Section Field Catalog")
    add_para(doc, "This is the implementation-level catalog for the API response. It is aligned with the attached Markdown examples and OpenAPI YAML.")
    add_table(doc, ["Module", "Field", "Type", "Aggregation", "Business Definition"], FIELD_CATALOG, widths=[1.25, 2.7, 1.25, 1.55, 4.0], font_size=6.8)

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(portrait)
    add_header_footer(doc)

    add_heading(doc, "11. Classification and Band Rules")
    add_table(doc, ["Domain", "Range", "Band"], BAND_RULES, widths=[2.1, 1.45, 3.8], font_size=7.8)

    add_heading(doc, "12. Error Contract")
    add_table(doc, ["HTTP", "Error Code", "Meaning", "Client Action"], ERRORS, widths=[0.65, 1.9, 2.6, 2.85], font_size=7.6)

    add_heading(doc, "13. Security, Refresh, and Observability")
    add_table(doc, ["Area", "Requirement"], [
        ("Authentication", "Use bearer JWT or enterprise gateway token. Reject missing or expired tokens with 401."),
        ("Authorization", "Validate tenant, hierarchy, region, branch, and agent scope before returning data."),
        ("PII", "Do not expose customer PII in summary APIs. Return counts, bands, and amounts unless a separate privileged drill-down API is approved."),
        ("Refresh", "Expose dailyRefreshStatus, latestTxnDate, latestCommissionDate, and partial-data warning when source refresh is delayed."),
        ("Caching", "Cache by agentCode, windowCodes, sectionCodes, filters, include flags, and referenceDataVersion."),
        ("Observability", "Return requestId and propagate it to gateway, aggregation service, model service, and data store logs."),
        ("Localization", "Future-ready narrative localization can use Accept-Language without changing the schema."),
    ], widths=[1.6, 5.9], font_size=7.8)

    add_heading(doc, "14. Frontend Consumption Guide")
    add_table(doc, ["Frontend Element", "API Source", "Display Rule"], [
        ("Dashboard cards", "sections.<section>.overall", "Use compact cards with value, delta, band, and narrative."),
        ("Trend charts", "sections.<section>.trendSummary.periods[]", "Use periodCode as stable x-axis key and periodLabel for display."),
        ("Breakdown filters", "sections.transactionSummary.breakdowns", "Render channel, transaction type, and service type filters from returned arrays."),
        ("Section drill-down", "GET /portfolio-summary/sections/{sectionCode}", "Use when a widget opens details for one module."),
        ("Chart drill-through", "GET /portfolio-trends", "Pass agentCode + sectionCodes + grain + period dates."),
        ("Recommendation panel", "analyticsReportingManagement.recommendations", "Group by module and severity; show prescriptive actions first."),
        ("Reference data", "GET /reference/portfolio-codes", "Use to configure valid window, grain, and section filters."),
    ], widths=[1.8, 2.9, 2.8], font_size=7.8)

    add_heading(doc, "15. Generated API Artifacts")
    add_table(doc, ["Artifact", "Path"], [
        ("Enhanced Word specification", str(OUT_DOCX)),
        ("Full Word specification with source examples", str(OUT_FULL_DOCX)),
        ("Enhanced Markdown specification", str(OUT_MD)),
        ("Enhanced OpenAPI YAML", str(OUT_YAML)),
        ("Copied source Markdown reference", str(OUT_SOURCE_MD)),
    ], widths=[2.2, 5.3], font_size=7.8, header_fill=TEAL)

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(doc.sections[-1])
    add_header_footer(doc)
    add_heading(doc, "16. Full API Specification and Examples From Attached Markdown")
    add_note(
        doc,
        "Full fetched content",
        "This section renders the complete attached Markdown API specification, including endpoint details, section examples, trend examples, drill-down contract, and error example. Code blocks are wrapped as normal paragraphs so Pages/Word can paginate them without hiding text.",
        fill=LIGHT_TEAL,
    )
    source_markdown = SOURCE_MD.read_text(encoding="utf-8") if SOURCE_MD.exists() else ""
    render_markdown_source(doc, source_markdown)

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(doc.sections[-1])
    add_header_footer(doc)
    add_heading(doc, "17. Full OpenAPI YAML Contract")
    add_note(
        doc,
        "Complete machine-readable API spec",
        "The full OpenAPI YAML is included below and also written as a standalone YAML file for import into API tools.",
        fill=LIGHT_BLUE,
    )
    add_code(doc, enhanced_yaml())

    add_header_footer(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    doc.save(OUT_FULL_DOCX)


def write_outputs() -> None:
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")
    OUT_YAML.write_text(enhanced_yaml(), encoding="utf-8")
    if SOURCE_MD.exists():
        shutil.copyfile(SOURCE_MD, OUT_SOURCE_MD)
    build_docx()


if __name__ == "__main__":
    write_outputs()
    print(OUT_DOCX)
    print(OUT_FULL_DOCX)
    print(OUT_MD)
    print(OUT_YAML)
