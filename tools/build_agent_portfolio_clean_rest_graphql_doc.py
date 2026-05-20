from __future__ import annotations

import json
import shutil
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCS = ROOT / "docs" / "Agent_Portfolio_Summary_API_Spec_REST_GraphQL_Clean.docx"
OUT_DESKTOP = Path("/Users/gireesha/Desktop/Agent_Portfolio_Summary_API_Spec_Full_With_Examples.docx")
OUT_MD = ROOT / "docs" / "Agent_Portfolio_Summary_API_Spec_REST_GraphQL_Clean.md"

NAVY = "F59E0B"
TEAL = "D97706"
SLATE = "334155"
BLUE_LIGHT = "FEF3C7"
GRAY_LIGHT = "FFFBEB"
WHITE = "111827"


MODULES = [
    ("Transaction Summary", "transactionSummary", "Transactions, enrollments, lead services, ease services, cash movement, channel/type/service breakdowns."),
    ("Commission Summary", "commissionSummary", "SBOS Autopay commission, eligible transactions, eligible amount, commission rank, average commission."),
    ("System Performance", "systemPerformance", "Success/failure ratio, business declines, technical declines, top BD/TD response code."),
    ("Target & Campaign Management", "targetCampaignManagement", "Targets assigned, achievement, pending count, remaining days, target status, daily required pace."),
    ("Anomaly Risk Management", "anomalyRiskManagement", "Suspicious transactions, suspicious customers, suspicious amount, anomaly cases, risk band."),
    ("Agent Audit Management", "agentAuditManagement", "Audits initiated, audit marks, ongoing audits, recent audit list, last audited date."),
    ("Analytics Reporting Management", "analyticsReportingManagement", "Descriptive, diagnostic, predictive, and prescriptive recommendations."),
    ("Enterprise Agent Performance Risk Management", "enterpriseAgentPerformanceRiskManagement", "EAPRM iScore, band, component scores, customers engaged, attendance percentage."),
]

WINDOWS = [
    ("LAST_7_DAYS", "DAY", "7 periods", "Daily chart and recent operational action."),
    ("LAST_3_MONTHS", "MONTH", "3 periods", "Recent month-wise performance comparison."),
    ("CURRENT_FY_QUARTERS", "QUARTER", "4 periods", "Q1-Q4 of the current financial year."),
    ("CUSTOM", "DAY / MONTH / QUARTER", "Based on request", "Custom analytics range using fromDate and toDate."),
]

ENDPOINTS = [
    ("GET", "/sbos-ibpm/v1/agents/{agentCode}/portfolio-summary", "Complete dashboard summary for selected windows."),
    ("GET", "/sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/{sectionCode}", "One section only, for drill-down screens."),
    ("GET", "/sbos-ibpm/v1/agents/{agentCode}/portfolio-trends", "Chart-ready trends by DAY, MONTH, or QUARTER."),
    ("POST", "/sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/query", "Advanced filtered summary request."),
    ("POST", "/graphql", "GraphQL endpoint for frontend field selection."),
]

SECTION_CODES = [
    ("TRANSACTION_SUMMARY", "transactionSummary"),
    ("COMMISSION_SUMMARY", "commissionSummary"),
    ("SYSTEM_PERFORMANCE", "systemPerformance"),
    ("TARGET_CAMPAIGN_MANAGEMENT", "targetCampaignManagement"),
    ("ANOMALY_RISK_MANAGEMENT", "anomalyRiskManagement"),
    ("AGENT_AUDIT_MANAGEMENT", "agentAuditManagement"),
    ("ANALYTICS_REPORTING_MANAGEMENT", "analyticsReportingManagement"),
    ("ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT", "enterpriseAgentPerformanceRiskManagement"),
]

MODULE_WISE_APIS = [
    ("Transaction Summary", "TRANSACTION_SUMMARY", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/TRANSACTION_SUMMARY"),
    ("Commission Summary", "COMMISSION_SUMMARY", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/COMMISSION_SUMMARY"),
    ("System Performance", "SYSTEM_PERFORMANCE", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/SYSTEM_PERFORMANCE"),
    ("Target & Campaign Management", "TARGET_CAMPAIGN_MANAGEMENT", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/TARGET_CAMPAIGN_MANAGEMENT"),
    ("Anomaly Risk Management", "ANOMALY_RISK_MANAGEMENT", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/ANOMALY_RISK_MANAGEMENT"),
    ("Agent Audit Management", "AGENT_AUDIT_MANAGEMENT", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/AGENT_AUDIT_MANAGEMENT"),
    ("Analytics Reporting Management", "ANALYTICS_REPORTING_MANAGEMENT", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/ANALYTICS_REPORTING_MANAGEMENT"),
    ("Enterprise Agent Performance Risk Management", "ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT"),
]

MODULE_HIGHLIGHTS = [
    (
        "Transaction Summary",
        "Daily total transactions, financial/non-financial split, amount involved, cash inflow/outflow, top channel, top transaction type, service mix.",
        "Month-wise transaction count, amount involved, financial transaction ratio, cash movement, enrollment/lead/ease service movement.",
        "Quarter-wise transaction growth, amount growth, channel mix, financial transaction ratio trend, service adoption trend.",
    ),
    (
        "Commission Summary",
        "Daily commission earned, eligible transaction count, eligible amount, average commission, rank movement.",
        "Month-wise commission earned, eligible transaction value, average commission, rank trend.",
        "Quarter-wise commission growth, rank movement, commission productivity, eligible transaction quality.",
    ),
    (
        "System Performance",
        "Daily success ratio, failure ratio, BD count, TD count, top BD/TD response code.",
        "Month-wise success/failure trend, dominant decline type, top recurring response code.",
        "Quarter-wise system health band, error movement, operational stability comparison.",
    ),
    (
        "Target & Campaign Management",
        "Daily target achievement, pending transaction count, required transactions per day, target status.",
        "Month-wise targets assigned, achieved, pending, average achievement percent.",
        "Quarter-wise target completion, campaign effectiveness, missed target concentration, achievement band.",
    ),
    (
        "Anomaly Risk Management",
        "Daily suspicious transaction count/ratio, suspicious amount, suspicious customers, top anomaly case.",
        "Month-wise suspicious amount, anomaly case count, repeated customers, risk band movement.",
        "Quarter-wise risk trend, fraud/anomaly concentration, high-risk pattern movement.",
    ),
    (
        "Agent Audit Management",
        "Daily audits initiated/closed, ongoing audits, audit marks when available, latest audit status.",
        "Month-wise audit volume, average marks, ongoing audit backlog, audit source split.",
        "Quarter-wise audit quality, recurring audit source, compliance improvement, critical audit count.",
    ),
    (
        "Analytics Reporting Management",
        "Daily recommendations by type/severity, open/accepted/closed recommendations, top module.",
        "Month-wise recommendation volume, major/critical recommendations, acceptance and closure trend.",
        "Quarter-wise recommendation effectiveness, unresolved recommendation load, priority trend.",
    ),
    (
        "Enterprise Agent Performance Risk Management",
        "Daily iScore, band, operational/compliance/financial health scores, attendance, customers engaged.",
        "Month-wise iScore movement, component score trend, attendance trend, customer engagement trend.",
        "Quarter-wise iScore band movement, component score improvement, risk-performance trajectory.",
    ),
]

RESPONSE_FIELDS = [
    ("requestId", "string", "Unique request id for support and logs."),
    ("schemaVersion", "string", "Schema version, for example 2.0."),
    ("referenceDataVersion", "string", "Version of channel, section, transaction, and service reference data."),
    ("generatedAt", "datetime", "Timestamp when response was generated."),
    ("agent", "object", "Agent identity, status, branch, region, hierarchy level."),
    ("portfolioStatus", "object", "Overall band, daily refresh status, freshness, warning."),
    ("timeWindows[]", "array", "One object per requested window."),
    ("timeWindows[].sections", "object", "Module sections requested by the caller."),
    ("recommendations[]", "array", "Cross-module action recommendations."),
]

SECTION_PATTERN = [
    ("overall", "Summary values for dashboard cards."),
    ("trendSummary.periods[]", "Chart and drill-down rows. Each row has periodCode, label, grain, fromDate, toDate."),
    ("breakdowns", "Optional channel/type/service/product split when includeBreakdowns=true."),
    ("narrative", "Plain-language explanation for the agent or supervisor."),
]

FIELD_CATALOG = [
    ("Transaction", "financialTxnCount, nonFinancialTxnCount, totalTxnCount", "Counts of financial, non-financial, and all transactions."),
    ("Transaction", "amountInvolved, cashInflow, cashOutflow", "Money totals using { value, currency }."),
    ("Transaction", "financialTxnRatio, averageFinancialTxnRatio", "Weighted percentage ratios."),
    ("Transaction", "breakdowns.byChannel / byTransactionType / byServiceType", "ONUS/OFFUS, AEPS/RUPAY/etc., CD/CW/FT split."),
    ("Commission", "totalCommissionEarned, eligibleTxnAmount", "Commission amount and value eligible for commission."),
    ("Commission", "totalEligibleTxnCount, commissionRank, averageCommissionEarned", "Eligibility count, rank movement, average earning."),
    ("System Performance", "successRatio, failureRatio", "Success and failure percentage."),
    ("System Performance", "businessDeclineCount, technicalDeclineCount", "BD/TD counts and top response codes."),
    ("TCM", "totalTargetsAssigned, totalTargetAchievement, targets[]", "Target progress and daily action requirement."),
    ("ARM", "suspiciousTxnRatio, suspiciousAmount, topAnomalyCase", "Risk volume, value, and dominant anomaly."),
    ("AAM", "totalAuditsInitiated, auditedMarksScored, ongoingAuditCount", "Audit status and score."),
    ("ADRM", "recommendations[]", "Descriptive, diagnostic, predictive, prescriptive actions."),
    ("EAPRM", "iScore, band, operationalEfficiencyScore, complianceScore", "Enterprise score and score components."),
    ("EAPRM", "financialHealthScore, totalCustomersEngaged, attendancePercentage", "Financial health, reach, and attendance."),
]

ERRORS = [
    ("400", "INVALID_REQUEST / INVALID_WINDOW_CODE", "Fix query/body and use reference endpoint."),
    ("401", "UNAUTHORIZED", "Re-authenticate."),
    ("403", "FORBIDDEN_AGENT_SCOPE", "Caller cannot view requested agent."),
    ("404", "AGENT_NOT_FOUND", "Verify agentCode."),
    ("409", "REFRESH_IN_PROGRESS", "Retry or show last successful snapshot."),
    ("422", "INVALID_FILTER", "Refresh reference data and retry."),
    ("500", "PORTFOLIO_SUMMARY_FAILED", "Log requestId and show friendly error."),
]

INDEX_ROWS = [
    ("1", "Introduction and purpose"),
    ("2", "Document index"),
    ("3", "Glossary"),
    ("4", "API surface"),
    ("5", "Complete API and module-wise APIs"),
    ("6", "Modules returned"),
    ("7", "Time windows"),
    ("8", "Section codes"),
    ("9", "Module-wise required highlights"),
    ("10", "Complete request and response schemas"),
    ("11", "Module-wise request and response schemas"),
    ("12", "Standard response shape"),
    ("13", "REST API specification"),
    ("14", "REST request and response examples"),
    ("15", "GraphQL API, query, variables, and response examples"),
    ("16", "Field catalog"),
    ("17", "Error contract"),
    ("18", "Frontend guidance"),
]

GLOSSARY = [
    ("Agent / BC", "Business Correspondent or agent whose portfolio is retrieved."),
    ("Portfolio Summary", "Consolidated agent view across transactions, commission, system health, targets, risk, audits, analytics, and EAPRM."),
    ("Window Code", "Named time range such as LAST_7_DAYS, LAST_3_MONTHS, or CURRENT_FY_QUARTERS."),
    ("Grain", "Trend period type: DAY, MONTH, or QUARTER."),
    ("Section Code", "API code identifying one module, such as TRANSACTION_SUMMARY."),
    ("Overall", "Summary values used by dashboard cards."),
    ("Trend Summary", "Array of period rows used by charts and drill-down."),
    ("Breakdowns", "Channel, transaction type, service type, product, or category splits."),
    ("Narrative", "Simple business explanation returned with KPI values."),
    ("EAPRM", "Enterprise Agent Performance Risk Management score and component score module."),
    ("iScore", "Enterprise score from EAPRM used to classify agent risk/performance."),
    ("BD / TD", "Business Decline and Technical Decline response code categories."),
]

COMPLETE_REQUEST_SCHEMA = [
    ("agentCode", "path string", "Yes", "BC00012345", "Agent/BC code in URL path."),
    ("windowCodes[]", "array enum", "No", "LAST_7_DAYS, LAST_3_MONTHS, CURRENT_FY_QUARTERS", "Time windows to return."),
    ("sectionCodes[]", "array enum", "No", "TRANSACTION_SUMMARY, COMMISSION_SUMMARY", "Modules to include. Omit for all modules."),
    ("includeTrends", "boolean", "No", "true", "Return trendSummary.periods[]."),
    ("includeBreakdowns", "boolean", "No", "true", "Return breakdowns object."),
    ("includeNarratives", "boolean", "No", "true", "Return narrative text."),
    ("filters.channelCodes[]", "array string", "No", "ONUS, OFFUS", "Filter by channel."),
    ("filters.transactionTypeCodes[]", "array string", "No", "AEPS, RUPAY", "Filter by transaction type."),
    ("filters.serviceTypeCodes[]", "array string", "No", "CD, CW, FT", "Filter by service type."),
    ("filters.enrollmentTypeCodes[]", "array string", "No", "PMJJBY, PMSBY, APY", "Filter by enrollment product."),
    ("filters.leadServiceTypeCodes[]", "array string", "No", "Loan, RD, FD", "Filter by lead service."),
    ("filters.easeServiceTypeCodes[]", "array string", "No", "Debit Card Hotlisting", "Filter by ease service."),
]

COMPLETE_RESPONSE_SCHEMA = [
    ("requestId", "string", "Yes", "Correlation id for support and logs."),
    ("schemaVersion", "string", "Yes", "API schema version."),
    ("referenceDataVersion", "string", "Yes", "Reference taxonomy version."),
    ("generatedAt", "datetime", "Yes", "Response timestamp."),
    ("agent", "object", "Yes", "Agent identity and hierarchy."),
    ("portfolioStatus", "object", "Yes", "Overall band, refresh state, data freshness."),
    ("timeWindows[]", "array", "Yes", "One object per requested window."),
    ("timeWindows[].sections", "object", "Yes", "Map of module JSON sections."),
    ("recommendations[]", "array", "No", "Cross-module recommendations."),
]

MODULE_REQUEST_SCHEMA = [
    ("agentCode", "path string", "Yes", "Agent/BC code."),
    ("sectionCode", "path enum", "Yes", "One value from Section Codes table."),
    ("windowCodes", "query CSV", "No", "LAST_7_DAYS,LAST_3_MONTHS,CURRENT_FY_QUARTERS."),
    ("includeTrends", "query boolean", "No", "Include day/month/quarter trend rows."),
    ("includeBreakdowns", "query boolean", "No", "Include split by channel/type/service where applicable."),
    ("includeNarratives", "query boolean", "No", "Include business explanation."),
]

MODULE_RESPONSE_SCHEMA = [
    ("requestId", "string", "Yes", "Correlation id."),
    ("generatedAt", "datetime", "Yes", "Response timestamp."),
    ("agent", "object", "Yes", "Agent identity."),
    ("sectionCode", "enum", "Yes", "Requested module code."),
    ("timeWindows[]", "array", "Yes", "Window rows for requested module only."),
    ("timeWindows[].section.overall", "object", "Yes", "Module dashboard summary."),
    ("timeWindows[].section.trendSummary.periods[]", "array", "If requested", "Day/month/quarter trend rows."),
    ("timeWindows[].section.breakdowns", "object", "If requested", "Module-specific breakdowns."),
    ("timeWindows[].section.narrative", "string", "If requested", "Module explanation."),
]

COMBINED_SCHEMA = {
    "completePortfolioRequest": {
        "agentCode": "path parameter",
        "body": {
            "windowCodes": ["LAST_7_DAYS", "LAST_3_MONTHS", "CURRENT_FY_QUARTERS"],
            "sectionCodes": ["TRANSACTION_SUMMARY", "COMMISSION_SUMMARY"],
            "includeTrends": True,
            "includeBreakdowns": True,
            "includeNarratives": True,
            "filters": {"channelCodes": ["ONUS", "OFFUS"]},
        },
    },
    "completePortfolioResponse": {
        "requestId": "string",
        "agent": "Agent",
        "portfolioStatus": "PortfolioStatus",
        "timeWindows": [{"code": "WindowCode", "grain": "TrendGrain", "sections": "PortfolioSections"}],
        "recommendations": ["Recommendation"],
    },
    "modulePortfolioRequest": {
        "agentCode": "path parameter",
        "sectionCode": "path parameter",
        "query": {"windowCodes": "CSV", "includeTrends": True, "includeBreakdowns": True},
    },
    "modulePortfolioResponse": {
        "requestId": "string",
        "sectionCode": "SectionCode",
        "timeWindows": [{"code": "WindowCode", "section": "RequestedSection"}],
    },
}

REST_ADVANCED_REQUEST = {
    "windowCodes": ["LAST_7_DAYS", "LAST_3_MONTHS", "CURRENT_FY_QUARTERS"],
    "sectionCodes": [
        "TRANSACTION_SUMMARY",
        "COMMISSION_SUMMARY",
        "SYSTEM_PERFORMANCE",
        "TARGET_CAMPAIGN_MANAGEMENT",
        "ANOMALY_RISK_MANAGEMENT",
        "AGENT_AUDIT_MANAGEMENT",
        "ANALYTICS_REPORTING_MANAGEMENT",
        "ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT",
    ],
    "includeTrends": True,
    "includeBreakdowns": True,
    "includeNarratives": True,
    "filters": {
        "channelCodes": ["ONUS", "OFFUS"],
        "transactionTypeCodes": ["AEPS", "RUPAY", "SHG", "TPD", "IMPS"],
        "serviceTypeCodes": ["CD", "CW", "FT"],
        "enrollmentTypeCodes": ["PMJJBY", "PMSBY", "APY"],
        "leadServiceTypeCodes": ["Loan", "RD", "FD"],
        "easeServiceTypeCodes": ["Debit Card Hotlisting", "Account Enquiry"],
    },
}

REST_FILLED_RESPONSE = {
    "requestId": "req_01HY9Z6N0W5D4Z7D6T7E8K9P2A",
    "schemaVersion": "2.0",
    "referenceDataVersion": "svc-taxonomy-2026-05",
    "generatedAt": "2026-05-20T09:30:00+05:30",
    "agent": {
        "agentCode": "BC00012345",
        "displayName": "Sample BC Agent",
        "status": "ACTIVE",
        "branchCode": "BR001",
        "regionCode": "RG-SOUTH",
        "hierarchyLevel": "AGENT",
    },
    "portfolioStatus": {
        "overallBand": "Stable",
        "dailyRefreshStatus": "COMPLETED",
        "dataFreshness": {
            "latestTxnDate": "2026-05-19",
            "latestCommissionDate": "2026-05-18",
            "latencyMinutes": 18,
        },
        "warning": None,
    },
    "timeWindows": [
        {
            "code": "LAST_7_DAYS",
            "label": "Last 7 Days",
            "fromDate": "2026-05-13",
            "toDate": "2026-05-19",
            "grain": "DAY",
            "periodStatus": "COMPLETED",
            "sections": {
                "transactionSummary": {
                    "overall": {
                        "financialTxnCount": 184,
                        "nonFinancialTxnCount": 59,
                        "enrollmentCount": 12,
                        "leadServiceCount": 8,
                        "easeServiceCount": 19,
                        "totalTxnCount": 243,
                        "amountInvolved": {"value": 1854000.0, "currency": "INR"},
                        "financialTxnRatio": 75.72,
                        "averageFinancialTxnRatio": 72.16,
                        "cashInflow": {"value": 984000.0, "currency": "INR"},
                        "cashOutflow": {"value": 870000.0, "currency": "INR"},
                    },
                    "trendSummary": {
                        "grain": "DAY",
                        "periods": [
                            {
                                "periodCode": "2026-05-13",
                                "periodLabel": "13 May 2026",
                                "grain": "DAY",
                                "fromDate": "2026-05-13",
                                "toDate": "2026-05-13",
                                "financialTxnCount": 22,
                                "nonFinancialTxnCount": 9,
                                "totalTxnCount": 31,
                                "amountInvolved": {"value": 212000.0, "currency": "INR"},
                            },
                            {
                                "periodCode": "2026-05-14",
                                "periodLabel": "14 May 2026",
                                "grain": "DAY",
                                "fromDate": "2026-05-14",
                                "toDate": "2026-05-14",
                                "financialTxnCount": 25,
                                "nonFinancialTxnCount": 9,
                                "totalTxnCount": 34,
                                "amountInvolved": {"value": 244000.0, "currency": "INR"},
                            },
                        ],
                    },
                    "breakdowns": {
                        "byChannel": [
                            {"code": "ONUS", "totalTxnCount": 138, "amountInvolved": 1042000.0},
                            {"code": "OFFUS", "totalTxnCount": 105, "amountInvolved": 812000.0},
                        ],
                        "byTransactionType": [
                            {"code": "AEPS", "totalTxnCount": 121, "amountInvolved": 1021000.0},
                            {"code": "RUPAY", "totalTxnCount": 44, "amountInvolved": 276000.0},
                        ],
                        "byServiceType": [
                            {"code": "CD", "totalTxnCount": 74, "amountInvolved": 984000.0},
                            {"code": "CW", "totalTxnCount": 68, "amountInvolved": 870000.0},
                        ],
                    },
                    "narrative": "Financial transactions formed 75.72% of total activity, above the agent average.",
                },
                "commissionSummary": {
                    "overall": {
                        "totalCommissionEarned": {"value": 12450.75, "currency": "INR"},
                        "totalEligibleTxnCount": 171,
                        "eligibleTxnAmount": {"value": 1619000.0, "currency": "INR"},
                        "commissionRank": {"rank": 18, "peerGroup": "BR001", "movement": 3},
                        "totalTxnCount": 243,
                        "averageCommissionEarned": {"value": 72.81, "currency": "INR"},
                    },
                    "trendSummary": {
                        "grain": "DAY",
                        "periods": [
                            {
                                "periodCode": "2026-05-13",
                                "commissionEarned": {"value": 1525.5, "currency": "INR"},
                                "eligibleTxnCount": 21,
                                "commissionRank": 21,
                            }
                        ],
                    },
                    "narrative": "Commission improved because eligible AEPS and cash deposit transactions increased.",
                },
                "systemPerformance": {
                    "overall": {
                        "failureRatio": 2.88,
                        "successRatio": 97.12,
                        "businessDeclineCount": 5,
                        "technicalDeclineCount": 2,
                        "topBusinessDeclineResponseCode": {"code": "BD01", "label": "Insufficient balance", "count": 3},
                        "topTechnicalDeclineResponseCode": {"code": "TD09", "label": "Issuer timeout", "count": 2},
                        "band": "Low",
                    },
                    "trendSummary": {"grain": "DAY", "periods": [{"periodCode": "2026-05-13", "successRatio": 96.77, "failureRatio": 3.23}]},
                    "narrative": "System health is stable; most declines are business declines, not technical errors.",
                },
                "targetCampaignManagement": {
                    "overall": {"totalTargetsAssigned": 3, "totalTargetAchievement": 68.4},
                    "targets": [
                        {
                            "targetId": "TGT-AEPS-CW-MAY",
                            "targetName": "AEPS Cash Withdrawal May Campaign",
                            "achievementPercent": 68.4,
                            "targetStatus": "Average",
                            "txnPerDayRequired": 10,
                            "pendingTxnCount": 126,
                            "remainingDays": 12,
                        }
                    ],
                    "trendSummary": {"grain": "DAY", "periods": [{"periodCode": "2026-05-13", "achievementPercent": 64.2, "targetStatus": "Average"}]},
                    "narrative": "Complete about 10 more AEPS cash withdrawal transactions per day to finish on time.",
                },
                "anomalyRiskManagement": {
                    "overall": {
                        "suspiciousTxnRatio": 1.65,
                        "totalSuspiciousTxnCount": 4,
                        "totalAnomalyCaseCount": 2,
                        "suspiciousCustomerCount": 3,
                        "suspiciousAmount": {"value": 44000.0, "currency": "INR"},
                        "topAnomalyCase": {"code": "ODD_TIME_TXN", "label": "Odd time transaction", "count": 2},
                        "riskBand": "Very Low Risk",
                    },
                    "trendSummary": {"grain": "DAY", "periods": [{"periodCode": "2026-05-13", "suspiciousTxnCount": 1, "riskBand": "Very Low Risk"}]},
                    "narrative": "Risk is low; the main anomaly pattern is odd-time transactions.",
                },
                "agentAuditManagement": {
                    "overall": {
                        "totalAuditsInitiated": 2,
                        "auditedMarksScored": 84.0,
                        "ongoingAuditCount": 1,
                        "lastAuditedOn": "2026-05-11",
                    },
                    "recentAudits": [
                        {
                            "auditId": "AUD-2026-0511-01",
                            "auditName": "ARM-triggered risk audit",
                            "auditSource": "ARM",
                            "auditType": "Risk Based",
                            "marksScored": 84.0,
                            "status": "CLOSED",
                        }
                    ],
                    "trendSummary": {"grain": "DAY", "periods": [{"periodCode": "2026-05-13", "ongoingAuditCount": 1}]},
                    "narrative": "Latest audit result is High; one audit remains open.",
                },
                "analyticsReportingManagement": {
                    "overall": {"totalRecommendations": 2, "majorRecommendationCount": 1, "minorRecommendationCount": 1},
                    "recommendations": [
                        {
                            "module": "TCM",
                            "type": "PRESCRIPTIVE",
                            "severity": "Major",
                            "message": "Increase AEPS cash withdrawal by 10 transactions per day.",
                        },
                        {
                            "module": "ARM",
                            "type": "DIAGNOSTIC",
                            "severity": "Minor",
                            "message": "Most suspicious activity is linked to odd-time transactions.",
                        },
                    ],
                    "trendSummary": {"grain": "DAY", "periods": [{"periodCode": "2026-05-13", "totalRecommendations": 1, "topModule": "TCM"}]},
                },
                "enterpriseAgentPerformanceRiskManagement": {
                    "overall": {
                        "iScore": 842,
                        "band": "Stable",
                        "operationalEfficiencyScore": 88.2,
                        "behaviouralEfficiencyScore": 81.5,
                        "complianceScore": 86.0,
                        "financialHealthScore": 79.7,
                        "totalCustomersEngaged": 96,
                        "attendancePercentage": 94.0,
                    },
                    "trendSummary": {"grain": "DAY", "periods": [{"periodCode": "2026-05-13", "iScore": 834, "band": "Stable"}]},
                    "narrative": "The agent is stable, with strong operations and good attendance.",
                },
            },
        },
        {
            "code": "LAST_3_MONTHS",
            "label": "Last 3 Months",
            "fromDate": "2026-03-01",
            "toDate": "2026-05-31",
            "grain": "MONTH",
            "periodStatus": "COMPLETED",
            "sections": {
                "transactionSummary": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "periodLabel": "Mar 2026", "totalTxnCount": 920, "amountInvolved": {"value": 6810000.0, "currency": "INR"}},
                        {"periodCode": "2026-04", "periodLabel": "Apr 2026", "totalTxnCount": 1015, "amountInvolved": {"value": 7425000.0, "currency": "INR"}},
                        {"periodCode": "2026-05", "periodLabel": "May 2026", "totalTxnCount": 832, "amountInvolved": {"value": 5980000.0, "currency": "INR"}},
                    ]},
                    "narrative": "Monthly transaction volume peaked in April and remains healthy in May month-to-date.",
                },
                "commissionSummary": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "commissionEarned": {"value": 43800.0, "currency": "INR"}, "commissionRank": 24},
                        {"periodCode": "2026-04", "commissionEarned": {"value": 51250.0, "currency": "INR"}, "commissionRank": 19},
                        {"periodCode": "2026-05", "commissionEarned": {"value": 39400.0, "currency": "INR"}, "commissionRank": 18},
                    ]},
                    "narrative": "Commission rank improved from 24 to 18 over the last three months.",
                },
                "systemPerformance": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "successRatio": 96.4, "failureRatio": 3.6, "band": "Moderate"},
                        {"periodCode": "2026-04", "successRatio": 97.0, "failureRatio": 3.0, "band": "Low"},
                        {"periodCode": "2026-05", "successRatio": 97.12, "failureRatio": 2.88, "band": "Low"},
                    ]},
                    "narrative": "System performance is improving month over month.",
                },
                "targetCampaignManagement": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "totalTargetAchievement": 61.2, "targetStatus": "Average"},
                        {"periodCode": "2026-04", "totalTargetAchievement": 66.7, "targetStatus": "Average"},
                        {"periodCode": "2026-05", "totalTargetAchievement": 68.4, "targetStatus": "Average"},
                    ]},
                    "narrative": "Target achievement is improving, but daily pace still needs attention.",
                },
                "anomalyRiskManagement": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "suspiciousTxnRatio": 2.4, "riskBand": "Very Low Risk"},
                        {"periodCode": "2026-04", "suspiciousTxnRatio": 1.9, "riskBand": "Very Low Risk"},
                        {"periodCode": "2026-05", "suspiciousTxnRatio": 1.65, "riskBand": "Very Low Risk"},
                    ]},
                    "narrative": "Risk ratio reduced across the last three months.",
                },
                "agentAuditManagement": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "auditsInitiated": 1, "averageMarksScored": 78.0},
                        {"periodCode": "2026-04", "auditsInitiated": 1, "averageMarksScored": 82.0},
                        {"periodCode": "2026-05", "auditsInitiated": 2, "averageMarksScored": 84.0},
                    ]},
                    "narrative": "Audit score improved to 84 in May.",
                },
                "analyticsReportingManagement": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "totalRecommendations": 5, "majorCount": 2},
                        {"periodCode": "2026-04", "totalRecommendations": 4, "majorCount": 1},
                        {"periodCode": "2026-05", "totalRecommendations": 2, "majorCount": 1},
                    ]},
                    "narrative": "Recommendation volume is reducing, showing better portfolio stability.",
                },
                "enterpriseAgentPerformanceRiskManagement": {
                    "trendSummary": {"grain": "MONTH", "periods": [
                        {"periodCode": "2026-03", "iScore": 812, "band": "Stable"},
                        {"periodCode": "2026-04", "iScore": 831, "band": "Stable"},
                        {"periodCode": "2026-05", "iScore": 842, "band": "Stable"},
                    ]},
                    "narrative": "iScore improved steadily across the last three months.",
                },
            },
        },
        {
            "code": "CURRENT_FY_QUARTERS",
            "label": "Current Financial Year Quarters",
            "fromDate": "2026-04-01",
            "toDate": "2027-03-31",
            "grain": "QUARTER",
            "periodStatus": "ACTIVE",
            "sections": {
                "transactionSummary": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "totalTxnCount": 2680, "amountInvolved": {"value": 19930000.0, "currency": "INR"}},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "totalTxnCount": 2920, "amountInvolved": {"value": 21850000.0, "currency": "INR"}},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "totalTxnCount": 3150, "amountInvolved": {"value": 23600000.0, "currency": "INR"}},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "totalTxnCount": 3360, "amountInvolved": {"value": 25100000.0, "currency": "INR"}},
                    ]},
                },
                "commissionSummary": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "commissionEarned": {"value": 134450.0, "currency": "INR"}},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "commissionEarned": {"value": 148000.0, "currency": "INR"}},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "commissionEarned": {"value": 159500.0, "currency": "INR"}},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "commissionEarned": {"value": 171200.0, "currency": "INR"}},
                    ]},
                },
                "systemPerformance": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "successRatio": 97.1, "band": "Low"},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "successRatio": 97.4, "band": "Low"},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "successRatio": 97.8, "band": "Very Low"},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "successRatio": 98.0, "band": "Very Low"},
                    ]},
                },
                "targetCampaignManagement": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "totalTargetAchievement": 68.4, "targetStatus": "Average"},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "totalTargetAchievement": 78.0, "targetStatus": "High"},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "totalTargetAchievement": 86.0, "targetStatus": "High"},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "totalTargetAchievement": 94.0, "targetStatus": "High"},
                    ]},
                },
                "anomalyRiskManagement": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "suspiciousTxnRatio": 1.65, "riskBand": "Very Low Risk"},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "suspiciousTxnRatio": 1.4, "riskBand": "Very Low Risk"},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "suspiciousTxnRatio": 1.2, "riskBand": "Very Low Risk"},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "suspiciousTxnRatio": 1.0, "riskBand": "Very Low Risk"},
                    ]},
                },
                "agentAuditManagement": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "auditsInitiated": 4, "averageMarksScored": 84.0},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "auditsInitiated": 3, "averageMarksScored": 86.0},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "auditsInitiated": 3, "averageMarksScored": 88.0},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "auditsInitiated": 2, "averageMarksScored": 90.0},
                    ]},
                },
                "analyticsReportingManagement": {
                    "trendSummary": {"grain": "QUARTER", "periods": [
                        {"periodCode": "FY2026-27-Q1", "periodStatus": "ACTIVE", "totalRecommendations": 11, "majorCount": 4},
                        {"periodCode": "FY2026-27-Q2", "periodStatus": "FORECAST", "totalRecommendations": 9, "majorCount": 3},
                        {"periodCode": "FY2026-27-Q3", "periodStatus": "FORECAST", "totalRecommendations": 7, "majorCount": 2},
                        {"periodCode": "FY2026-27-Q4", "periodStatus": "FORECAST", "totalRecommendations": 5, "majorCount": 1},
                    ]},
                },
                "enterpriseAgentPerformanceRiskManagement": {
                    "trendSummary": {
                        "grain": "QUARTER",
                        "periods": [
                            {
                                "periodCode": "FY2026-27-Q1",
                                "periodLabel": "Q1 FY 2026-27",
                                "fromDate": "2026-04-01",
                                "toDate": "2026-06-30",
                                "periodStatus": "ACTIVE",
                                "iScore": 842,
                                "band": "Stable",
                            },
                            {
                                "periodCode": "FY2026-27-Q2",
                                "periodLabel": "Q2 FY 2026-27",
                                "fromDate": "2026-07-01",
                                "toDate": "2026-09-30",
                                "periodStatus": "FORECAST",
                                "iScore": 855,
                                "band": "Stable",
                            },
                            {
                                "periodCode": "FY2026-27-Q3",
                                "periodLabel": "Q3 FY 2026-27",
                                "fromDate": "2026-10-01",
                                "toDate": "2026-12-31",
                                "periodStatus": "FORECAST",
                                "iScore": 872,
                                "band": "Stable",
                            },
                            {
                                "periodCode": "FY2026-27-Q4",
                                "periodLabel": "Q4 FY 2026-27",
                                "fromDate": "2027-01-01",
                                "toDate": "2027-03-31",
                                "periodStatus": "FORECAST",
                                "iScore": 890,
                                "band": "Stable",
                            },
                        ],
                    }
                }
            },
        },
    ],
    "recommendations": [
        {
            "module": "TCM",
            "type": "PRESCRIPTIVE",
            "priority": "Major",
            "message": "Prioritize AEPS cash withdrawal activity for the next 12 days.",
        }
    ],
}

GRAPHQL_SCHEMA = """type Query {
  agentPortfolioSummary(input: AgentPortfolioSummaryInput!): AgentPortfolioSummary!
  agentPortfolioSection(input: AgentPortfolioSectionInput!): PortfolioSectionResponse!
  agentPortfolioTrends(input: AgentPortfolioTrendInput!): AgentPortfolioTrendResponse!
}

input AgentPortfolioSummaryInput {
  agentCode: ID!
  windowCodes: [WindowCode!] = [LAST_7_DAYS, LAST_3_MONTHS, CURRENT_FY_QUARTERS]
  sectionCodes: [SectionCode!]
  includeTrends: Boolean = true
  includeBreakdowns: Boolean = true
  includeNarratives: Boolean = true
  filters: PortfolioFiltersInput
}

enum WindowCode { LAST_7_DAYS LAST_3_MONTHS CURRENT_FY_QUARTERS CUSTOM }
enum TrendGrain { DAY MONTH QUARTER }
enum SectionCode {
  TRANSACTION_SUMMARY
  COMMISSION_SUMMARY
  SYSTEM_PERFORMANCE
  TARGET_CAMPAIGN_MANAGEMENT
  ANOMALY_RISK_MANAGEMENT
  AGENT_AUDIT_MANAGEMENT
  ANALYTICS_REPORTING_MANAGEMENT
  ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT
}

type AgentPortfolioSummary {
  requestId: ID!
  schemaVersion: String!
  generatedAt: DateTime!
  agent: Agent!
  portfolioStatus: PortfolioStatus!
  timeWindows: [TimeWindowSummary!]!
  recommendations: [Recommendation!]!
}

type TimeWindowSummary {
  code: WindowCode!
  label: String!
  fromDate: Date!
  toDate: Date!
  grain: TrendGrain!
  sections: PortfolioSections!
}

type PortfolioSections {
  transactionSummary: TransactionSummarySection
  commissionSummary: CommissionSummarySection
  systemPerformance: SystemPerformanceSection
  targetCampaignManagement: TargetCampaignManagementSection
  anomalyRiskManagement: AnomalyRiskManagementSection
  agentAuditManagement: AgentAuditManagementSection
  analyticsReportingManagement: AnalyticsReportingManagementSection
  enterpriseAgentPerformanceRiskManagement: EAPRMSection
}"""

GRAPHQL_QUERY = """query AgentPortfolioSummary($input: AgentPortfolioSummaryInput!) {
  agentPortfolioSummary(input: $input) {
    requestId
    generatedAt
    agent { agentCode displayName status branchCode regionCode }
    portfolioStatus { overallBand dailyRefreshStatus }
    timeWindows {
      code
      label
      grain
      sections {
        transactionSummary {
          overall {
            totalTxnCount
            financialTxnRatio
            amountInvolved { value currency }
          }
          trendSummary {
            periods { periodCode periodLabel totalTxnCount }
          }
          narrative
        }
        commissionSummary {
          overall {
            totalCommissionEarned { value currency }
            commissionRank { rank peerGroup movement }
          }
        }
        enterpriseAgentPerformanceRiskManagement {
          overall { iScore band attendancePercentage }
          narrative
        }
      }
    }
    recommendations { module type priority message }
  }
}"""

GRAPHQL_VARIABLES = {
    "input": {
        "agentCode": "BC00012345",
        "windowCodes": ["LAST_7_DAYS", "LAST_3_MONTHS", "CURRENT_FY_QUARTERS"],
        "sectionCodes": ["TRANSACTION_SUMMARY", "COMMISSION_SUMMARY", "ENTERPRISE_AGENT_PERFORMANCE_RISK_MANAGEMENT"],
        "includeTrends": True,
        "includeBreakdowns": True,
        "includeNarratives": True,
        "filters": {
            "channelCodes": ["ONUS", "OFFUS"],
            "transactionTypeCodes": ["AEPS", "RUPAY"],
            "serviceTypeCodes": ["CD", "CW"],
        },
    }
}

GRAPHQL_RESPONSE = {
    "data": {
        "agentPortfolioSummary": {
            "requestId": "req_01HY9Z6N0W5D4Z7D6T7E8K9P2A",
            "generatedAt": "2026-05-20T09:30:00+05:30",
            "agent": {"agentCode": "BC00012345", "displayName": "Sample BC Agent", "status": "ACTIVE", "branchCode": "BR001", "regionCode": "RG-SOUTH"},
            "portfolioStatus": {"overallBand": "Stable", "dailyRefreshStatus": "COMPLETED"},
            "timeWindows": [
                {
                    "code": "LAST_7_DAYS",
                    "label": "Last 7 Days",
                    "grain": "DAY",
                    "sections": {
                        "transactionSummary": {
                            "overall": {"totalTxnCount": 243, "financialTxnRatio": 75.72, "amountInvolved": {"value": 1854000.0, "currency": "INR"}},
                            "trendSummary": {"periods": [{"periodCode": "2026-05-13", "periodLabel": "13 May 2026", "totalTxnCount": 31}]},
                            "narrative": "Financial transactions formed 75.72% of total activity.",
                        },
                        "commissionSummary": {
                            "overall": {"totalCommissionEarned": {"value": 12450.75, "currency": "INR"}, "commissionRank": {"rank": 18, "peerGroup": "BR001", "movement": 3}}
                        },
                        "enterpriseAgentPerformanceRiskManagement": {
                            "overall": {"iScore": 842, "band": "Stable", "attendancePercentage": 94.0},
                            "narrative": "The agent is stable, with strong operations and good attendance.",
                        },
                    },
                }
            ],
            "recommendations": [{"module": "TCM", "type": "PRESCRIPTIVE", "priority": "Major", "message": "Prioritize AEPS cash withdrawal activity."}],
        }
    }
}


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=48, start=60, bottom=48, end=60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 7.6) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[tuple | list], widths: list[float] | None = None, font_size: float = 7.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, NAVY)
        set_cell_text(cell, header, bold=True, color=WHITE, size=font_size)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            set_cell_text(cells[col_index], value, size=font_size)
            if row_index % 2 == 1:
                shade_cell(cells[col_index], "F8FAFC")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)


def add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def set_footer_with_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix = footer.add_run("Agent Portfolio Summary API | REST + GraphQL | Page ")
        prefix.font.size = Pt(7)
        prefix.font.color.rgb = RGBColor.from_string(SLATE)
        add_field_run(footer, "PAGE")
        mid = footer.add_run(" of ")
        mid.font.size = Pt(7)
        mid.font.color.rgb = RGBColor.from_string(SLATE)
        add_field_run(footer, "NUMPAGES")


def shade_paragraph(paragraph, fill: str = GRAY_LIGHT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def wrap_code(code: str, width: int = 94) -> list[str]:
    wrapped: list[str] = []
    for raw in code.strip().splitlines():
        line = raw.rstrip()
        if len(line) <= width:
            wrapped.append(line)
            continue
        indent = line[: len(line) - len(line.lstrip(" "))]
        wrapped.extend(textwrap.wrap(line, width=width, subsequent_indent=indent + "  ", break_long_words=False, break_on_hyphens=False))
    return wrapped


def add_code(doc: Document, code: str) -> None:
    for line in wrap_code(code):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.84
        p.paragraph_format.left_indent = Inches(0.05)
        shade_paragraph(p)
        run = p.add_run(line or " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(5.8)
        run.font.color.rgb = RGBColor.from_string("111827")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(4 if level == 1 else 2)
    p.paragraph_format.space_after = Pt(1)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.line_spacing = 1.0
    p.add_run(text)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, BLUE_LIGHT)
    set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    p.add_run(body)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.42)
    section.right_margin = Inches(0.42)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(8.4)
    for name, size, color in [("Heading 1", 13, NAVY), ("Heading 2", 10.5, TEAL), ("Heading 3", 9.2, SLATE)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    set_footer_with_page_numbers(doc)


def json_block(value: object) -> str:
    return json.dumps(value, indent=2, ensure_ascii=False)


def module_section_response(section_code: str, json_field: str) -> dict:
    windows = []
    for window in REST_FILLED_RESPONSE["timeWindows"]:
        section_data = window["sections"].get(json_field)
        if not section_data:
            continue
        windows.append(
            {
                "code": window["code"],
                "label": window["label"],
                "fromDate": window["fromDate"],
                "toDate": window["toDate"],
                "grain": window["grain"],
                "section": section_data,
            }
        )
    return {
        "requestId": "req_module_01HY9Z6N0W5D4Z7D6T7E8K9P2A",
        "generatedAt": REST_FILLED_RESPONSE["generatedAt"],
        "agent": REST_FILLED_RESPONSE["agent"],
        "sectionCode": section_code,
        "timeWindows": windows,
    }


def build_markdown() -> str:
    module_api_rows = "\n".join(f"| {module} | `{code}` | `{endpoint}` |" for module, code, endpoint in MODULE_WISE_APIS)
    highlight_rows = "\n".join(
        f"| {module} | {daily} | {monthly} | {quarterly} |"
        for module, daily, monthly, quarterly in MODULE_HIGHLIGHTS
    )
    return "\n\n".join(
        [
            "# Agent Portfolio Summary API Specification - REST + GraphQL",
            "## Index\n| Section | Topic |\n|---|---|\n" + "\n".join(f"| {section} | {topic} |" for section, topic in INDEX_ROWS),
            "## Glossary\n| Term | Meaning |\n|---|---|\n" + "\n".join(f"| {term} | {meaning} |" for term, meaning in GLOSSARY),
            "## Module-wise APIs\n| Module | Section Code | REST API |\n|---|---|---|\n" + module_api_rows,
            "## Module-wise Required Highlights\n| Module | Last 7 Days - Daily | Last 3 Months - Monthly | Current FY 4 Quarters - Quarterly |\n|---|---|---|---|\n" + highlight_rows,
            "## Complete Request Schema\n| Field | Type | Required | Example | Meaning |\n|---|---|---|---|---|\n" + "\n".join(f"| `{field}` | {type_} | {required} | {example} | {meaning} |" for field, type_, required, example, meaning in COMPLETE_REQUEST_SCHEMA),
            "## Complete Response Schema\n| Field | Type | Required | Meaning |\n|---|---|---|---|\n" + "\n".join(f"| `{field}` | {type_} | {required} | {meaning} |" for field, type_, required, meaning in COMPLETE_RESPONSE_SCHEMA),
            "## Module Request Schema\n| Field | Type | Required | Meaning |\n|---|---|---|---|\n" + "\n".join(f"| `{field}` | {type_} | {required} | {meaning} |" for field, type_, required, meaning in MODULE_REQUEST_SCHEMA),
            "## Module Response Schema\n| Field | Type | Required | Meaning |\n|---|---|---|---|\n" + "\n".join(f"| `{field}` | {type_} | {required} | {meaning} |" for field, type_, required, meaning in MODULE_RESPONSE_SCHEMA),
            "## REST Advanced Query Request\n```json\n" + json_block(REST_ADVANCED_REQUEST) + "\n```",
            "## REST Filled Response\n```json\n" + json_block(REST_FILLED_RESPONSE) + "\n```",
            "## GraphQL Schema\n```graphql\n" + GRAPHQL_SCHEMA + "\n```",
            "## GraphQL Query\n```graphql\n" + GRAPHQL_QUERY + "\n```",
            "## GraphQL Variables\n```json\n" + json_block(GRAPHQL_VARIABLES) + "\n```",
            "## GraphQL Filled Response\n```json\n" + json_block(GRAPHQL_RESPONSE) + "\n```",
        ]
    )


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Agent Portfolio Summary API Specification")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    r = subtitle.add_run("Yellow Theme Template | REST + GraphQL | Complete and Module-wise Schemas with Examples")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    add_table(
        doc,
        ["Document Field", "Value"],
        [
            ("Version", "4.0 yellow theme"),
            ("Prepared on", "20 May 2026"),
            ("Base path", "/sbos-ibpm/v1"),
            ("Primary key", "agentCode"),
            ("Scope", "Index, glossary, REST, GraphQL, complete schemas, module-wise schemas, filled examples, pagination"),
        ],
        widths=[1.5, 6.2],
        font_size=7.8,
    )

    add_heading(doc, "1. Introduction")
    add_para(doc, "This API gives a complete portfolio summary for a Business Correspondent / Agent. A dashboard can show transactions, commission, system health, targets, risk, audits, recommendations, and enterprise iScore from one response.")
    add_para(doc, "The response is designed for both technical users and business users: numeric KPIs are returned together with trend rows, breakdowns, bands, and simple narrative explanations.")
    add_note(doc, "Template", "This version uses a compact yellow theme, tight spacing, wrapped code examples, and footer pagination.")

    add_heading(doc, "2. Document Index")
    add_table(doc, ["Section", "Topic"], INDEX_ROWS, widths=[0.85, 6.9], font_size=7.4)

    add_heading(doc, "3. Glossary")
    add_table(doc, ["Term", "Meaning"], GLOSSARY, widths=[1.75, 6.0], font_size=7.2)

    add_heading(doc, "4. API Surface")
    add_table(doc, ["Method", "Endpoint", "Purpose"], ENDPOINTS, widths=[0.55, 3.2, 4.0], font_size=7.2)

    add_heading(doc, "5. Complete API and Module-wise API")
    add_table(
        doc,
        ["API Type", "Endpoint", "When To Use"],
        [
            ("Complete REST API", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary", "Load full dashboard with all or selected modules."),
            ("Complete REST Query API", "POST /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/query", "Use advanced filters and explicit module/window selection."),
            ("Module-wise REST API", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-summary/sections/{sectionCode}", "Load or refresh one module independently."),
            ("Trend API", "GET /sbos-ibpm/v1/agents/{agentCode}/portfolio-trends", "Load chart-only data by section and grain."),
            ("GraphQL API", "POST /graphql", "Fetch only fields needed by the frontend widget."),
        ],
        widths=[1.5, 3.4, 2.9],
        font_size=7.1,
    )
    add_table(doc, ["Module", "Section Code", "Module-wise REST API"], MODULE_WISE_APIS, widths=[1.95, 2.25, 3.6], font_size=6.6)

    add_heading(doc, "6. Modules Returned")
    add_table(doc, ["Module", "JSON Field", "What It Means"], MODULES, widths=[2.0, 2.2, 3.6], font_size=7.0)

    add_heading(doc, "7. Time Windows")
    add_table(doc, ["Window", "Grain", "Count", "Usage"], WINDOWS, widths=[1.45, 1.45, 1.2, 3.7], font_size=7.2)

    add_heading(doc, "8. Section Codes")
    add_table(doc, ["Section Code", "JSON Field"], SECTION_CODES, widths=[3.5, 3.5], font_size=7.2)

    add_heading(doc, "9. Module-wise Required Highlights")
    add_note(
        doc,
        "Mandatory trend coverage",
        "Every module must support day-wise highlights for LAST_7_DAYS, month-wise highlights for LAST_3_MONTHS, and quarter-wise highlights for all four quarters of the current financial year.",
    )
    add_table(
        doc,
        ["Module", "Last 7 Days - Daily", "Last 3 Months - Monthly", "Current FY 4 Quarters - Quarterly"],
        MODULE_HIGHLIGHTS,
        widths=[1.45, 2.1, 2.1, 2.15],
        font_size=5.9,
    )

    add_heading(doc, "10. Complete Request and Response Schemas")
    add_table(doc, ["Request Field", "Type", "Required", "Example", "Meaning"], COMPLETE_REQUEST_SCHEMA, widths=[1.7, 1.15, 0.75, 1.7, 2.5], font_size=6.6)
    add_table(doc, ["Response Field", "Type", "Required", "Meaning"], COMPLETE_RESPONSE_SCHEMA, widths=[2.2, 1.2, 0.75, 3.6], font_size=7.0)
    add_heading(doc, "10.1 Combined Complete API Schema", 2)
    add_code(doc, json_block(COMBINED_SCHEMA))

    add_heading(doc, "11. Module-wise Request and Response Schemas")
    add_table(doc, ["Module Request Field", "Type", "Required", "Meaning"], MODULE_REQUEST_SCHEMA, widths=[2.0, 1.4, 0.8, 3.6], font_size=7.0)
    add_table(doc, ["Module Response Field", "Type", "Required", "Meaning"], MODULE_RESPONSE_SCHEMA, widths=[2.5, 1.2, 0.8, 3.3], font_size=7.0)

    add_heading(doc, "12. Standard Response Shape")
    add_table(doc, ["Field", "Type", "Meaning"], RESPONSE_FIELDS, widths=[2.2, 1.3, 4.3], font_size=7.2)
    add_table(doc, ["Every Section Contains", "Purpose"], SECTION_PATTERN, widths=[2.3, 5.4], font_size=7.2)

    add_heading(doc, "13. REST API Specification")
    add_table(
        doc,
        ["Endpoint", "Required", "Optional"],
        [
            ("GET /agents/{agentCode}/portfolio-summary", "agentCode", "windowCodes, includeTrends, includeBreakdowns, includeNarratives"),
            ("POST /agents/{agentCode}/portfolio-summary/query", "agentCode, request body", "filters, sectionCodes, windowCodes"),
            ("GET /agents/{agentCode}/portfolio-trends", "agentCode, grain", "sectionCodes, fromDate, toDate"),
        ],
        widths=[3.0, 2.0, 2.8],
        font_size=7.1,
    )

    add_heading(doc, "14. REST Example - Complete API Advanced Request")
    add_code(doc, "POST /sbos-ibpm/v1/agents/BC00012345/portfolio-summary/query\nAuthorization: Bearer <token>\nContent-Type: application/json\n\n" + json_block(REST_ADVANCED_REQUEST))

    add_heading(doc, "15. REST Example - Module-wise API Request")
    add_code(
        doc,
        "GET /sbos-ibpm/v1/agents/BC00012345/portfolio-summary/sections/TRANSACTION_SUMMARY"
        "?windowCodes=LAST_7_DAYS,LAST_3_MONTHS,CURRENT_FY_QUARTERS"
        "&includeTrends=true&includeBreakdowns=true&includeNarratives=true\n"
        "Authorization: Bearer <token>",
    )

    add_heading(doc, "16. REST Examples - Module-wise Requests and Responses")
    for section_code, json_field in SECTION_CODES:
        add_heading(doc, section_code, level=3)
        add_code(
            doc,
            f"GET /sbos-ibpm/v1/agents/BC00012345/portfolio-summary/sections/{section_code}"
            "?windowCodes=LAST_7_DAYS,LAST_3_MONTHS,CURRENT_FY_QUARTERS"
            "&includeTrends=true&includeBreakdowns=true&includeNarratives=true\n"
            "Authorization: Bearer <token>\n\n"
            + json_block(module_section_response(section_code, json_field)),
        )

    add_heading(doc, "17. REST Example - Filled Complete Response")
    add_note(doc, "Filled response", "This example contains real-looking values for all portfolio sections, plus daily, monthly, and quarterly trend windows.")
    add_code(doc, json_block(REST_FILLED_RESPONSE))

    add_heading(doc, "18. GraphQL API Specification")
    add_para(doc, "GraphQL is useful when the frontend wants only selected widgets or fields. It should call the same portfolio aggregation logic as REST.")
    add_code(doc, GRAPHQL_SCHEMA)

    add_heading(doc, "19. GraphQL Example - Query")
    add_code(doc, GRAPHQL_QUERY)

    add_heading(doc, "20. GraphQL Example - Variables")
    add_code(doc, json_block(GRAPHQL_VARIABLES))

    add_heading(doc, "21. GraphQL Example - Filled Response")
    add_code(doc, json_block(GRAPHQL_RESPONSE))

    add_heading(doc, "22. Field Catalog")
    add_table(doc, ["Module", "Field(s)", "Meaning"], FIELD_CATALOG, widths=[1.6, 2.8, 3.4], font_size=6.8)

    add_heading(doc, "23. Error Contract")
    add_table(doc, ["HTTP", "Code", "Client Action"], ERRORS, widths=[0.65, 2.35, 4.8], font_size=7.1)

    add_heading(doc, "24. Frontend Guidance")
    add_table(
        doc,
        ["Screen Area", "Use This Field", "Display Guidance"],
        [
            ("KPI cards", "sections.<module>.overall", "Show current value, band, and narrative."),
            ("Charts", "trendSummary.periods[]", "Use periodCode for x-axis and periodLabel for display."),
            ("Filters", "breakdowns", "Render channel/type/service splits only when returned."),
            ("Recommendations", "analyticsReportingManagement.recommendations", "Show prescriptive actions first, then diagnostic insights."),
            ("Drill-down", "section endpoint or trend endpoint", "Pass agentCode, sectionCode, grain, and period range."),
        ],
        widths=[1.5, 2.7, 3.6],
        font_size=7.1,
    )

    OUT_DOCS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCS)
    shutil.copyfile(OUT_DOCS, OUT_DESKTOP)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    print(OUT_DOCS)
    print(OUT_DESKTOP)
    print(OUT_MD)
