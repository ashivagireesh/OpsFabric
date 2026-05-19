from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "Blueprnt_Agent_Portfolio_API_Spec.docx"
OUT_MD = ROOT / "docs" / "Blueprnt_Agent_Portfolio_API_Spec.md"
OUT_REST_REQUEST = ROOT / "docs" / "Blueprnt_Agent_Portfolio_REST_Request.json"
OUT_REST_RESPONSE = ROOT / "docs" / "Blueprnt_Agent_Portfolio_REST_Response.json"
OUT_OPENAPI = ROOT / "docs" / "Blueprnt_Agent_Portfolio_OpenAPI.yaml"
OUT_GRAPHQL = ROOT / "docs" / "Blueprnt_Agent_Portfolio_GraphQL_Schema.graphql"

NAVY = "1F4E79"
TEAL = "0F766E"
SLATE = "334155"
LIGHT_BLUE = "DDEBF7"
LIGHT_TEAL = "DDF3F0"
LIGHT_GRAY = "F3F6F8"
MID_GRAY = "E2E8F0"
WHITE = "FFFFFF"


DOCUMENT_META = [
    ("Document title", "Blueprnt Agent Portfolio API Specification"),
    ("Version", "2.0"),
    ("Status", "Enterprise draft for solution, API, frontend, and data engineering review"),
    ("Prepared on", "19 May 2026"),
    ("Prepared for", "Agent Portfolio, SBOS Autopay, iBPM, and EAPRM integration"),
    ("Source reference", "Blueprnt API (2).docx"),
    ("Primary lookup key", "agentCode"),
    ("Primary consumers", "Agent-facing web app, mobile app, dashboard, BFF/API gateway, analytics services"),
]

MODULES = [
    (
        "AUTOPAY",
        "SBOS Autopay",
        "Smart Banking Operations and Integration Suite - Autopay",
        "Automated commission management for Business Correspondent timely commission calculation and payout transparency.",
        "Commission cards, eligibility trend, commission rank, projected commission narrative.",
    ),
    (
        "TCM",
        "Target and Campaign Management",
        "iBPM sub-module",
        "Configure criteria, assign targets, monitor BC/vendor performance, and trigger SLA-driven workflows.",
        "Target progress cards, daily target pace, pending transaction count, achievement band.",
    ),
    (
        "ARM",
        "Anomaly and Risk Management",
        "iBPM sub-module",
        "ML-powered anomaly detection, fraud identification, suspicious customer concentration, and transaction risk classification.",
        "Risk badge, suspicious transaction trend, top anomaly case, escalation panel.",
    ),
    (
        "SPM",
        "System Performance Management",
        "iBPM sub-module",
        "Near real-time and historical monitoring of transaction health, success, decline, and response-code patterns.",
        "Success/failure trend, top BD/TD response codes, operational health band.",
    ),
    (
        "AAM",
        "Agent Audit Management",
        "iBPM sub-module",
        "Agent audit lifecycle management, risk-based audit ranking, source-triggered audits, and outcome tracking.",
        "Audit timeline, marks scored, open audit count, last audit status.",
    ),
    (
        "ADRM",
        "Analytics Dashboard and Reporting Management",
        "iBPM sub-module",
        "Interactive analytics, custom reports, and descriptive, diagnostic, predictive, and prescriptive recommendations.",
        "Recommendation feed, severity badge, module-wise insights.",
    ),
    (
        "EAPRM",
        "Enterprise Agent Performance and Risk Management",
        "Enterprise AI/ML governance layer",
        "360-degree Agent Governance Framework combining operational efficiency, fraud risk, compliance, financial health, and behavior into an iScore.",
        "iScore hero card, score components, trend movement, agent performance band.",
    ),
]

TIME_WINDOWS = [
    (
        "LAST_7_DAYS",
        "Rolling daily series",
        "DAY",
        "2026-05-13",
        "2026-05-19",
        "Seven calendar dates ending on asOfDate. Always return all seven dates, with zero-filled missing days.",
    ),
    (
        "CURRENT_MONTH",
        "Month-to-date",
        "DAY or MONTH",
        "2026-05-01",
        "2026-05-19",
        "Current calendar month from the first day through asOfDate.",
    ),
    (
        "PREVIOUS_MONTH_1",
        "Previous full month",
        "MONTH",
        "2026-04-01",
        "2026-04-30",
        "Complete month immediately before the current month.",
    ),
    (
        "PREVIOUS_MONTH_2",
        "Second previous full month",
        "MONTH",
        "2026-03-01",
        "2026-03-31",
        "Complete month two months before the current month.",
    ),
    (
        "CURRENT_QUARTER",
        "Financial quarter-to-date",
        "DAY, WEEK, or MONTH",
        "2026-04-01",
        "2026-05-19",
        "Quarter containing asOfDate. For India FY 2026-27 this is Q1, Apr-Jun.",
    ),
    (
        "CURRENT_FY",
        "Financial year-to-date",
        "MONTH",
        "2026-04-01",
        "2026-05-19",
        "Financial year from configured FY start through asOfDate. Default FY start is April 1.",
    ),
    (
        "FY_QUARTER_COMPARISON",
        "Financial year quarter comparison",
        "QUARTER",
        "2026-04-01",
        "2027-03-31",
        "Return Q1, Q2, Q3, and Q4 buckets. Future quarters carry periodStatus NOT_STARTED unless forecast fields are requested.",
    ),
]

GLOSSARY = [
    ("Agent / BC", "Business Correspondent or field agent whose portfolio is displayed."),
    ("Agent Portfolio", "Single consolidated view of agent transactions, commission, performance, targets, risk, audits, analytics, and iScore."),
    ("SBOS Autopay", "Smart Banking Operations and Integration Suite component for commission calculation and commission transparency."),
    ("iBPM", "Intelligent Business Performance Management for FI-BC ecosystem operations, compliance, health supervision, and fraud mitigation."),
    ("EAPRM", "Enterprise Agent Performance and Risk Management platform for risk-performance scoring and governance."),
    ("iScore", "Enterprise score from 0 to 1000 generated from operational, behavioral, compliance, and financial-health dimensions."),
    ("ONUS / OFFUS", "Channel grouping for same-bank and other-bank transaction processing."),
    ("AEPS", "Aadhaar Enabled Payment System transaction type."),
    ("RUPAY", "RuPay card transaction type."),
    ("SHG", "Self Help Group related transaction type."),
    ("TPD", "Third-party deposit or partner-defined transaction category, configurable by bank."),
    ("IMPS", "Immediate Payment Service transaction type."),
    ("CD / CW / FT", "Cash Deposit, Cash Withdrawal, and Fund Transfer service types."),
    ("BD / TD", "Business Decline and Technical Decline response categories."),
    ("Narrative", "Plain-language explanation derived from KPI values, trends, bands, and recommended next action."),
    ("Financial Year", "Configurable accounting period. This document uses India FY by default: April 1 to March 31."),
]

REQUEST_FIELDS = [
    ("agentCode", "string", "Yes", "BC00012345", "Unique code for the agent/BC portfolio to retrieve."),
    ("asOfDate", "date", "No", "2026-05-19", "Date for computing current windows. Defaults to current business date."),
    ("financialYear", "string", "No", "2026-27", "Financial year label. If omitted, derived from asOfDate and FY start month."),
    ("financialYearStartMonth", "integer", "No", "4", "Configured FY start month. Default 4 for April."),
    ("timeWindows", "array<enum>", "Yes", "LAST_7_DAYS, CURRENT_FY", "List of required historical windows."),
    ("filters.channels", "array<string>", "No", "ONUS, OFFUS", "Channel filters. Values should come from configurable reference data."),
    ("filters.transactionTypes", "array<string>", "No", "AEPS, RUPAY, SHG, TPD, IMPS", "Transaction type filters."),
    ("filters.serviceTypes", "array<string>", "No", "CD, CW, FT", "Financial service type filters."),
    ("filters.enrollmentTypes", "array<string>", "No", "PMJJBY, PMSBY, APY", "Enrollment product filters."),
    ("filters.leadServiceTypes", "array<string>", "No", "Loan, RD, FD", "Lead service filters."),
    ("filters.easeServiceTypes", "array<string>", "No", "Debit Card Hotlisting, Account Enquiry", "Ease service filters."),
    ("includeNarratives", "boolean", "No", "true", "When true, return simple explanation text beside KPI values."),
    ("includeBenchmarks", "boolean", "No", "true", "When true, return peer/branch/regional benchmark values where available."),
    ("includeFrontendHints", "boolean", "No", "true", "When true, return display hints such as severity, chart type, and unit metadata."),
]

ROOT_RESPONSE_FIELDS = [
    ("requestId", "string", "Trace id for API observability and support."),
    ("generatedAt", "datetime", "Server timestamp when response was generated."),
    ("agent", "object", "Agent identity, hierarchy, status, and display fields."),
    ("portfolioStatus", "object", "Overall band, refresh status, data freshness, and warning state."),
    ("timeWindows", "array<object>", "One object per requested time window. Each object contains all module sections."),
    ("recommendations", "array<object>", "Cross-module recommendation feed for frontend display."),
    ("referenceDataVersion", "string", "Version of transaction/service/channel taxonomy used in computation."),
    ("schemaVersion", "string", "API response schema version."),
]

METRIC_CATALOG = [
    ("Transaction", "Financial transactions", "financialTxnCount", "SUM", "Count", "Total financial transactions for selected filters."),
    ("Transaction", "Non-financial transactions", "nonFinancialTxnCount", "SUM", "Count", "Total non-financial transactions."),
    ("Transaction", "Enrollments", "enrollmentCount", "SUM", "Count", "PMJJBY, PMSBY, APY, and future enrollment types."),
    ("Transaction", "Lead services", "leadServiceCount", "SUM", "Count", "Loan, RD, FD, and future lead service types."),
    ("Transaction", "Ease services", "easeServiceCount", "SUM", "Count", "Debit card hotlisting, account enquiry, and future ease services."),
    ("Transaction", "Total transactions", "totalTxnCount", "SUM", "Count", "Financial plus non-financial transaction activity."),
    ("Transaction", "Amount involved", "amountInvolved", "SUM", "Money", "Gross transaction amount for included financial flows."),
    ("Transaction", "Financial transaction ratio", "financialTxnRatio", "Weighted ratio", "Percent", "financialTxnCount / totalTxnCount * 100."),
    ("Transaction", "Average financial txn ratio", "averageFinancialTxnRatio", "Weighted average", "Percent", "Peer/agent historical comparison ratio."),
    ("Transaction", "Cash inflow", "cashInflow", "SUM", "Money", "Cash deposits and inflow-equivalent services."),
    ("Transaction", "Cash outflow", "cashOutflow", "SUM", "Money", "Cash withdrawals and outflow-equivalent services."),
    ("AUTOPAY", "Total commission earned", "totalCommissionEarned", "SUM", "Money", "Commission accrued or paid for selected period."),
    ("AUTOPAY", "Eligible transactions", "totalEligibleTxnCount", "SUM", "Count", "Commission-eligible transactions."),
    ("AUTOPAY", "Eligible transaction amount", "eligibleTxnAmount", "SUM", "Money", "Amount involved in commission-eligible transactions."),
    ("AUTOPAY", "Commission rank", "commissionRank", "Latest plus movement", "Rank", "Agent rank in configured peer group."),
    ("AUTOPAY", "Total transactions", "totalTxnCount", "SUM", "Count", "Total transactions considered by commission rules."),
    ("AUTOPAY", "Average commission earned", "averageCommissionEarned", "Weighted average", "Money", "totalCommissionEarned / eligibleTxnCount."),
    ("SPM", "Failure ratio", "failureRatio", "Weighted ratio", "Percent", "Failed transactions / total transactions * 100."),
    ("SPM", "Success ratio", "successRatio", "Weighted ratio", "Percent", "Successful transactions / total transactions * 100."),
    ("SPM", "Business declines", "businessDeclineCount", "SUM", "Count", "Total BD response outcomes."),
    ("SPM", "Technical declines", "technicalDeclineCount", "SUM", "Count", "Total TD response outcomes."),
    ("SPM", "Top BD response code", "topBusinessDeclineResponseCode", "Mode", "Code", "Most frequent business decline code."),
    ("SPM", "Top TD response code", "topTechnicalDeclineResponseCode", "Mode", "Code", "Most frequent technical decline code."),
    ("TCM", "Targets assigned", "totalTargetsAssigned", "SUM", "Count", "Total active targets assigned to the agent."),
    ("TCM", "Target achievement", "totalTargetAchievement", "Weighted ratio", "Percent", "Overall target achievement percentage."),
    ("TCM", "Target-wise achievement", "targetWiseAchievementPercent", "Per target ratio", "Percent", "Achievement per target/campaign."),
    ("TCM", "Required transactions per day", "txnPerDayRequired", "Derived", "Count/day", "pendingTxnCount / remainingDays."),
    ("TCM", "Target status", "targetStatus", "Band rule", "Band", "Not Started, Critical, Low, Average, High, Achieved."),
    ("TCM", "Pending transactions", "pendingTxnCount", "Derived", "Count", "Target count minus achieved count."),
    ("TCM", "Remaining days", "remainingDays", "Date diff", "Days", "Campaign end date minus asOfDate."),
    ("ARM", "Suspicious transaction ratio", "suspiciousTxnRatio", "Weighted ratio", "Percent", "Suspicious transactions / total transactions * 100."),
    ("ARM", "Suspicious transactions", "totalSuspiciousTxnCount", "SUM", "Count", "Total flagged suspicious transactions."),
    ("ARM", "Anomaly cases involved", "totalAnomalyCaseCount", "SUM", "Count", "Distinct anomaly cases linked to the agent."),
    ("ARM", "Suspicious customers", "suspiciousCustomerCount", "Distinct count", "Count", "Distinct suspicious customers associated with agent."),
    ("ARM", "Suspicious amount", "suspiciousAmount", "SUM", "Money", "Amount involved in suspicious transactions."),
    ("ARM", "Top anomaly case", "topAnomalyCase", "Mode by severity/count", "Case", "Most frequent or most severe anomaly case type."),
    ("AAM", "Audits initiated", "totalAuditsInitiated", "SUM", "Count", "Total audits initiated for agent."),
    ("AAM", "Audited marks scored", "auditedMarksScored", "Latest or weighted average", "Percent/Score", "Marks for completed audits."),
    ("AAM", "Ongoing audits", "ongoingAuditCount", "Current count", "Count", "Audits not yet closed."),
    ("AAM", "Recent audit details", "recentAudits", "Latest N", "List", "Audit name, source, marks, type, status, and date."),
    ("AAM", "Last audited on", "lastAuditedOn", "MAX date", "Date", "Most recent completed or initiated audit date."),
    ("ADRM", "Descriptive recommendation", "descriptiveRecommendation", "Rule/AI generated", "Text", "What happened."),
    ("ADRM", "Diagnostic recommendation", "diagnosticRecommendation", "Rule/AI generated", "Text", "Why it happened."),
    ("ADRM", "Predictive recommendation", "predictiveRecommendation", "Model output", "Text", "What is likely to happen."),
    ("ADRM", "Prescriptive recommendation", "prescriptiveRecommendation", "Rule/AI generated", "Text", "What action should be taken."),
    ("EAPRM", "iScore", "iScore", "Latest score", "0-1000", "Enterprise performance and risk score."),
    ("EAPRM", "Operational efficiency", "operationalEfficiencyScore", "Latest score", "0-100", "Transaction productivity and system outcome quality."),
    ("EAPRM", "Behavioral efficiency", "behaviouralEfficiencyScore", "Latest score", "0-100", "Behavioral pattern and field conduct indicator."),
    ("EAPRM", "Compliance score", "complianceScore", "Latest score", "0-100", "Compliance and audit adherence indicator."),
    ("EAPRM", "Financial health score", "financialHealthScore", "Latest score", "0-100", "Financial stability and commission health indicator."),
    ("EAPRM", "Customers engaged", "totalCustomersEngaged", "Distinct count", "Count", "Distinct customers served or engaged."),
    ("EAPRM", "Attendance percentage", "attendancePercentage", "Percent", "Percent", "Attendance as percentage rather than raw days."),
]

CLASSIFICATIONS = [
    ("TCM Achievement", "0%", "Not Started", "No action started. Trigger onboarding/reminder workflow."),
    ("TCM Achievement", "1-24%", "Critical", "Immediate intervention and daily target guidance."),
    ("TCM Achievement", "25-50%", "Low", "Below required pace. Recommend campaign acceleration."),
    ("TCM Achievement", "51-75%", "Average", "Recoverable. Track remaining days and required transactions per day."),
    ("TCM Achievement", "76-99%", "High", "Close to completion. Show final target gap."),
    ("TCM Achievement", "100% or more", "Achieved", "Target fulfilled. Show completion and rank movement."),
    ("ARM Risk", "0%", "Genuine", "No detected risk."),
    ("ARM Risk", "1-25%", "Very Low Risk", "Monitor only."),
    ("ARM Risk", "26-50%", "Medium Risk", "Review pattern and customer concentration."),
    ("ARM Risk", "51-75%", "High Risk", "Escalate to supervisor review."),
    ("ARM Risk", "76-99%", "Very High Risk", "Urgent investigation and enhanced monitoring."),
    ("ARM Risk", "100%", "Fraud", "Confirmed fraud. Trigger immediate action workflow."),
    ("SPM Error Rate", "0%", "Perfect", "No error impact."),
    ("SPM Error Rate", ">0-1%", "Very Low", "Nominal impact."),
    ("SPM Error Rate", ">1-3%", "Low", "Minor degradation."),
    ("SPM Error Rate", ">3-5%", "Moderate", "Visible business impact."),
    ("SPM Error Rate", ">5-10%", "High", "Major operational concern."),
    ("SPM Error Rate", ">10%", "Critical", "Immediate operational intervention."),
    ("AAM Audit Score", "100%", "Perfect", "No material audit weakness."),
    ("AAM Audit Score", "75-99%", "High", "Strong audit outcome."),
    ("AAM Audit Score", "50-74%", "Moderate", "Improvement required."),
    ("AAM Audit Score", "25-49%", "Low", "Weak control quality."),
    ("AAM Audit Score", "1-24%", "Very Low", "Severe corrective action needed."),
    ("AAM Audit Score", "0%", "Critical", "Severe breach and possible suspension risk."),
    ("ADRM Severity", "0%", "Informational", "Awareness only."),
    ("ADRM Severity", "1-25%", "Minor", "Low-priority alert."),
    ("ADRM Severity", "26-50%", "Moderate", "Track and respond."),
    ("ADRM Severity", "51-75%", "Major", "Escalation needed."),
    ("ADRM Severity", "76-100%", "High", "Immediate stakeholder intervention."),
    ("EAPRM iScore", "<600", "Critical", "Highest governance risk."),
    ("EAPRM iScore", "601-700", "High Risk", "Strong remediation required."),
    ("EAPRM iScore", "701-800", "Medium Risk", "Managed monitoring."),
    ("EAPRM iScore", "801-900", "Stable", "Acceptable performance."),
    ("EAPRM iScore", "901-950", "Good", "Strong performer."),
    ("EAPRM iScore", "951-1000", "Very Good", "Model agent performance."),
]

ERRORS = [
    ("400", "INVALID_REQUEST", "Malformed JSON, missing required field, or invalid date.", "Fix request payload."),
    ("401", "UNAUTHORIZED", "Missing, expired, or invalid token.", "Re-authenticate user or service."),
    ("403", "FORBIDDEN_AGENT_SCOPE", "Caller is not allowed to view this agent.", "Show access denied and log requestId."),
    ("404", "AGENT_NOT_FOUND", "Agent code does not exist in active or historical master.", "Ask user to verify agent code."),
    ("409", "REFRESH_IN_PROGRESS", "Daily portfolio refresh is running.", "Retry with backoff or display previous snapshot."),
    ("422", "INVALID_FILTER", "Filter is not configured for tenant/reference data version.", "Refresh reference data and retry."),
    ("429", "RATE_LIMITED", "Caller exceeded configured API limit.", "Retry after response header time."),
    ("500", "PORTFOLIO_COMPUTE_FAILED", "Unexpected aggregation or downstream failure.", "Show friendly message and report requestId."),
]

REST_RESPONSE_FIELD_MAP = [
    ("requestId", "string", "Yes", "Trace id generated by the API service."),
    ("schemaVersion", "string", "Yes", "Response schema version."),
    ("referenceDataVersion", "string", "Yes", "Reference data taxonomy version used for channel, transaction, and service mappings."),
    ("generatedAt", "datetime", "Yes", "Response generation timestamp."),
    ("agent.agentCode", "string", "Yes", "Agent or BC unique code."),
    ("agent.displayName", "string", "Yes", "Display name of the agent."),
    ("agent.status", "string", "Yes", "ACTIVE, INACTIVE, SUSPENDED, CLOSED, or tenant-defined status."),
    ("agent.branchCode", "string", "No", "Branch mapped to the agent."),
    ("agent.regionCode", "string", "No", "Region mapped to the agent."),
    ("agent.hierarchyLevel", "string", "No", "Hierarchy level of the returned portfolio subject."),
    ("portfolioStatus.overallBand", "string", "Yes", "Overall agent status band."),
    ("portfolioStatus.dailyRefreshStatus", "string", "Yes", "COMPLETED, RUNNING, FAILED, PARTIAL, or SKIPPED."),
    ("portfolioStatus.dataFreshness.latestTxnDate", "date", "No", "Latest transaction date included."),
    ("portfolioStatus.dataFreshness.latestCommissionDate", "date", "No", "Latest commission date included."),
    ("portfolioStatus.dataFreshness.latencyMinutes", "integer", "No", "Freshness lag in minutes."),
    ("portfolioStatus.warning", "object|null", "No", "Warning returned for stale, partial, or unavailable module data."),
    ("timeWindows[]", "array<object>", "Yes", "One object per requested time range."),
    ("recommendations[]", "array<object>", "No", "Cross-module recommendation feed."),
]

MODULE_RESPONSE_FIELD_MAP = [
    ("Time Window", "code", "enum", "LAST_7_DAYS, CURRENT_MONTH, PREVIOUS_MONTH_1, PREVIOUS_MONTH_2, CURRENT_QUARTER, CURRENT_FY, FY_QUARTER_COMPARISON."),
    ("Time Window", "label", "string", "Frontend display label."),
    ("Time Window", "fromDate / toDate", "date", "Inclusive date range represented by the bucket."),
    ("Time Window", "grain", "enum", "DAY, WEEK, MONTH, QUARTER, or FY."),
    ("Time Window", "periodStatus", "enum", "COMPLETED, ACTIVE, PARTIAL, NOT_STARTED, UNAVAILABLE."),
    ("transactionSummary", "financialTxnCount", "integer", "Financial transaction count."),
    ("transactionSummary", "nonFinancialTxnCount", "integer", "Non-financial transaction count."),
    ("transactionSummary", "enrollmentCount", "integer", "Enrollment count for PMJJBY, PMSBY, APY, and future enrollment types."),
    ("transactionSummary", "leadServiceCount", "integer", "Lead service count for Loan, RD, FD, and future lead service types."),
    ("transactionSummary", "easeServiceCount", "integer", "Ease service count for debit card hotlisting, account enquiry, and future ease services."),
    ("transactionSummary", "totalTxnCount", "integer", "Total transaction count."),
    ("transactionSummary", "amountInvolved", "Money", "Gross transaction amount."),
    ("transactionSummary", "financialTxnRatio", "decimal", "Financial transaction count divided by total transaction count."),
    ("transactionSummary", "averageFinancialTxnRatio", "decimal", "Agent or peer average financial transaction ratio."),
    ("transactionSummary", "cashInflow", "Money", "Cash-in amount."),
    ("transactionSummary", "cashOutflow", "Money", "Cash-out amount."),
    ("transactionSummary", "breakdowns.byChannel[]", "array", "Channel-wise metrics for ONUS, OFFUS, and future channels."),
    ("transactionSummary", "breakdowns.byTransactionType[]", "array", "Transaction type-wise metrics for AEPS, RUPAY, SHG, TPD, IMPS, and future types."),
    ("transactionSummary", "breakdowns.byServiceType[]", "array", "Service-wise metrics for CD, CW, FT, and future services."),
    ("transactionSummary", "series[]", "array", "Daily, monthly, or quarterly chart series depending on requested grain."),
    ("transactionSummary", "narrative", "string", "Agent-friendly explanation."),
    ("commissionSummary", "totalCommissionEarned", "Money", "Commission earned/accrued."),
    ("commissionSummary", "totalEligibleTxnCount", "integer", "Commission eligible transaction count."),
    ("commissionSummary", "eligibleTxnAmount", "Money", "Value of eligible transactions."),
    ("commissionSummary", "commissionRank", "object", "Rank, peer group, and movement."),
    ("commissionSummary", "totalTxnCount", "integer", "Total transactions considered by commission engine."),
    ("commissionSummary", "averageCommissionEarned", "Money", "Average commission per eligible transaction."),
    ("systemPerformance", "failureRatio", "decimal", "Failure percentage."),
    ("systemPerformance", "successRatio", "decimal", "Success percentage."),
    ("systemPerformance", "businessDeclineCount", "integer", "BD count."),
    ("systemPerformance", "technicalDeclineCount", "integer", "TD count."),
    ("systemPerformance", "topBusinessDeclineResponseCode", "object", "Most frequent BD response code."),
    ("systemPerformance", "topTechnicalDeclineResponseCode", "object", "Most frequent TD response code."),
    ("targetCampaignManagement", "totalTargetsAssigned", "integer", "Assigned target count."),
    ("targetCampaignManagement", "totalTargetAchievement", "decimal", "Overall target achievement percent."),
    ("targetCampaignManagement", "targets[]", "array", "Target-wise achievement, status, required pace, pending count, remaining days."),
    ("anomalyRiskManagement", "suspiciousTxnRatio", "decimal", "Suspicious transaction percentage."),
    ("anomalyRiskManagement", "totalSuspiciousTxnCount", "integer", "Suspicious transaction count."),
    ("anomalyRiskManagement", "totalAnomalyCaseCount", "integer", "Anomaly cases linked to agent."),
    ("anomalyRiskManagement", "suspiciousCustomerCount", "integer", "Distinct suspicious customer count."),
    ("anomalyRiskManagement", "suspiciousAmount", "Money", "Amount involved in suspicious transactions."),
    ("anomalyRiskManagement", "topAnomalyCase", "object", "Most frequent or severe anomaly case."),
    ("agentAuditManagement", "totalAuditsInitiated", "integer", "Audits initiated."),
    ("agentAuditManagement", "auditedMarksScored", "decimal", "Latest or weighted audit marks."),
    ("agentAuditManagement", "ongoingAuditCount", "integer", "Open audit count."),
    ("agentAuditManagement", "lastAuditedOn", "date", "Last audit date."),
    ("agentAuditManagement", "recentAudits[]", "array", "Recent audit name, source, score, type, and status."),
    ("analyticsReportingManagement", "recommendations[]", "array", "Module-wise descriptive, diagnostic, predictive, and prescriptive recommendations."),
    ("enterpriseAgentPerformanceRiskManagement", "iScore", "integer", "EAPRM iScore from 0 to 1000."),
    ("enterpriseAgentPerformanceRiskManagement", "band", "string", "Critical, High Risk, Medium Risk, Stable, Good, Very Good."),
    ("enterpriseAgentPerformanceRiskManagement", "operationalEfficiencyScore", "decimal", "Operational efficiency component score."),
    ("enterpriseAgentPerformanceRiskManagement", "behaviouralEfficiencyScore", "decimal", "Behavioral efficiency component score."),
    ("enterpriseAgentPerformanceRiskManagement", "complianceScore", "decimal", "Compliance score."),
    ("enterpriseAgentPerformanceRiskManagement", "financialHealthScore", "decimal", "Financial health score."),
    ("enterpriseAgentPerformanceRiskManagement", "totalCustomersEngaged", "integer", "Distinct customers engaged."),
    ("enterpriseAgentPerformanceRiskManagement", "attendancePercentage", "decimal", "Attendance percentage."),
]

INDEX_TERMS = [
    ("ADRM", "Analytics recommendations; see Sections 5, 10, 11, 15"),
    ("AEPS", "Transaction type filter and target examples; see Sections 7, 8, 11"),
    ("Agent Code", "Primary lookup key; see Sections 1, 8, 10"),
    ("Agent Portfolio", "Canonical response contract; see Sections 1, 5, 8, 9, 10"),
    ("AAM", "Agent audit management; see Sections 5, 10, 11, 12"),
    ("ARM", "Anomaly and risk management; see Sections 5, 10, 11, 12"),
    ("AUTOPAY", "Commission management; see Sections 5, 8, 10, 11"),
    ("Classification Bands", "Target, risk, error, audit, severity, and iScore bands; see Section 12"),
    ("Commission", "SBOS Autopay fields and rules; see Sections 8, 10, 11"),
    ("Daily Refresh", "Snapshot and data freshness contract; see Section 14.3"),
    ("EAPRM", "Enterprise iScore and component scores; see Sections 5, 10, 11, 12"),
    ("Financial Year", "FY windows and quarter comparison; see Section 7"),
    ("Frontend", "Integration flow and widget mapping; see Sections 6 and 15"),
    ("GraphQL", "Query, input, and schema contract; see Section 9"),
    ("iBPM", "TCM, ARM, SPM, AAM, ADRM module group; see Section 5"),
    ("iScore", "EAPRM score band and components; see Sections 10, 11, 12"),
    ("Narrative", "Plain-language explanation rule; see Sections 1, 13, 15"),
    ("REST", "Endpoint, request, response, and OpenAPI contract; see Sections 8 and Appendix A"),
    ("SPM", "System performance management; see Sections 5, 10, 11, 12"),
    ("TCM", "Target and campaign management; see Sections 5, 10, 11, 12"),
    ("Time Windows", "Last 7 days, months, quarter, FY, FY quarters; see Section 7"),
]

REST_REQUEST = """{
  "agentCode": "BC00012345",
  "asOfDate": "2026-05-19",
  "financialYear": "2026-27",
  "financialYearStartMonth": 4,
  "timeWindows": [
    "LAST_7_DAYS",
    "CURRENT_MONTH",
    "PREVIOUS_MONTH_1",
    "PREVIOUS_MONTH_2",
    "CURRENT_QUARTER",
    "CURRENT_FY",
    "FY_QUARTER_COMPARISON"
  ],
  "filters": {
    "channels": ["ONUS", "OFFUS"],
    "transactionTypes": ["AEPS", "RUPAY", "SHG", "TPD", "IMPS"],
    "serviceTypes": ["CD", "CW", "FT"],
    "enrollmentTypes": ["PMJJBY", "PMSBY", "APY"],
    "leadServiceTypes": ["Loan", "RD", "FD"],
    "easeServiceTypes": ["Debit Card Hotlisting", "Account Enquiry"]
  },
  "includeNarratives": true,
  "includeBenchmarks": true,
  "includeFrontendHints": true
}"""

REST_RESPONSE = """{
  "requestId": "req_01HY9Z6N0W5D4Z7D6T7E8K9P2A",
  "schemaVersion": "2.0",
  "referenceDataVersion": "svc-taxonomy-2026-05",
  "generatedAt": "2026-05-19T18:45:30+05:30",
  "agent": {
    "agentCode": "BC00012345",
    "displayName": "Sample BC Agent",
    "status": "ACTIVE",
    "branchCode": "BR001",
    "regionCode": "RG-SOUTH",
    "hierarchyLevel": "AGENT"
  },
  "portfolioStatus": {
    "overallBand": "Stable",
    "dailyRefreshStatus": "COMPLETED",
    "dataFreshness": {
      "latestTxnDate": "2026-05-19",
      "latestCommissionDate": "2026-05-18",
      "latencyMinutes": 18
    },
    "warning": null
  },
  "timeWindows": [
    {
      "code": "LAST_7_DAYS",
      "label": "Last 7 Days",
      "fromDate": "2026-05-13",
      "toDate": "2026-05-19",
      "grain": "DAY",
      "periodStatus": "COMPLETED",
      "transactionSummary": {
        "financialTxnCount": 184,
        "nonFinancialTxnCount": 59,
        "enrollmentCount": 12,
        "leadServiceCount": 8,
        "easeServiceCount": 19,
        "totalTxnCount": 243,
        "amountInvolved": { "value": 1854000.00, "currency": "INR" },
        "financialTxnRatio": 75.72,
        "averageFinancialTxnRatio": 72.16,
        "cashInflow": { "value": 984000.00, "currency": "INR" },
        "cashOutflow": { "value": 870000.00, "currency": "INR" },
        "breakdowns": {
          "byChannel": [
            { "code": "ONUS", "totalTxnCount": 138, "amountInvolved": 1042000.00 },
            { "code": "OFFUS", "totalTxnCount": 105, "amountInvolved": 812000.00 }
          ],
          "byTransactionType": [
            { "code": "AEPS", "totalTxnCount": 121, "amountInvolved": 1021000.00 },
            { "code": "RUPAY", "totalTxnCount": 44, "amountInvolved": 276000.00 }
          ],
          "byServiceType": [
            { "code": "CD", "totalTxnCount": 74, "amountInvolved": 984000.00 },
            { "code": "CW", "totalTxnCount": 68, "amountInvolved": 870000.00 }
          ]
        },
        "series": [
          { "date": "2026-05-13", "financialTxnCount": 22, "totalTxnCount": 31, "amountInvolved": 212000.00 },
          { "date": "2026-05-14", "financialTxnCount": 25, "totalTxnCount": 34, "amountInvolved": 244000.00 }
        ],
        "narrative": "Financial transactions formed 75.72% of total activity, which is above the agent's seven-day average."
      },
      "commissionSummary": {
        "totalCommissionEarned": { "value": 12450.75, "currency": "INR" },
        "totalEligibleTxnCount": 171,
        "eligibleTxnAmount": { "value": 1619000.00, "currency": "INR" },
        "commissionRank": { "rank": 18, "peerGroup": "BR001", "movement": 3 },
        "totalTxnCount": 243,
        "averageCommissionEarned": { "value": 72.81, "currency": "INR" },
        "narrative": "Commission is improving because eligible AEPS and CD transactions increased during the last seven days."
      },
      "systemPerformance": {
        "failureRatio": 2.88,
        "successRatio": 97.12,
        "businessDeclineCount": 5,
        "technicalDeclineCount": 2,
        "topBusinessDeclineResponseCode": { "code": "BD01", "label": "Insufficient balance", "count": 3 },
        "topTechnicalDeclineResponseCode": { "code": "TD09", "label": "Issuer timeout", "count": 2 },
        "band": "Low",
        "narrative": "System performance is healthy. Most declines are business declines rather than technical failures."
      },
      "targetCampaignManagement": {
        "totalTargetsAssigned": 3,
        "totalTargetAchievement": 68.40,
        "targets": [
          {
            "targetId": "TGT-AEPS-CW-MAY",
            "targetName": "AEPS Cash Withdrawal May Campaign",
            "achievementPercent": 68.40,
            "targetStatus": "Average",
            "txnPerDayRequired": 10,
            "pendingTxnCount": 126,
            "remainingDays": 12
          }
        ],
        "narrative": "The agent should complete about 10 more AEPS cash withdrawal transactions per day to finish the target on time."
      },
      "anomalyRiskManagement": {
        "suspiciousTxnRatio": 1.65,
        "totalSuspiciousTxnCount": 4,
        "totalAnomalyCaseCount": 2,
        "suspiciousCustomerCount": 3,
        "suspiciousAmount": { "value": 44000.00, "currency": "INR" },
        "topAnomalyCase": { "code": "ODD_TIME_TXN", "label": "Odd time transaction", "count": 2 },
        "riskBand": "Very Low Risk",
        "narrative": "Risk is low. The main anomaly pattern is odd-time transaction activity."
      },
      "agentAuditManagement": {
        "totalAuditsInitiated": 2,
        "auditedMarksScored": 84.00,
        "ongoingAuditCount": 1,
        "lastAuditedOn": "2026-05-11",
        "recentAudits": [
          {
            "auditId": "AUD-2026-0511-01",
            "auditName": "ARM-triggered risk audit",
            "auditSource": "ARM",
            "auditType": "Risk Based",
            "marksScored": 84.00,
            "status": "CLOSED"
          }
        ],
        "narrative": "Latest audit result is High. One audit remains open."
      },
      "analyticsReportingManagement": {
        "recommendations": [
          {
            "module": "TCM",
            "type": "PRESCRIPTIVE",
            "severity": "Major",
            "message": "Increase AEPS cash withdrawal by 10 transactions per day to meet the May target."
          },
          {
            "module": "ARM",
            "type": "DIAGNOSTIC",
            "severity": "Minor",
            "message": "Most suspicious activity is linked to odd-time transaction behavior."
          }
        ]
      },
      "enterpriseAgentPerformanceRiskManagement": {
        "iScore": 842,
        "band": "Stable",
        "operationalEfficiencyScore": 88.20,
        "behaviouralEfficiencyScore": 81.50,
        "complianceScore": 86.00,
        "financialHealthScore": 79.70,
        "totalCustomersEngaged": 96,
        "attendancePercentage": 94.00,
        "narrative": "The agent is stable, with strong operational efficiency and good attendance."
      }
    }
  ],
  "recommendations": [
    {
      "module": "TCM",
      "type": "PRESCRIPTIVE",
      "priority": "Major",
      "message": "Prioritize AEPS cash withdrawal activity for the next 12 days."
    }
  ]
}"""

OPENAPI_EXCERPT = """openapi: 3.1.0
info:
  title: Blueprnt Agent Portfolio API
  version: "2.0"
paths:
  /api/v1/agent-portfolios/query:
    post:
      summary: Retrieve complete agent portfolio by agent code
      operationId: queryAgentPortfolio
      tags: [Agent Portfolio]
      security:
        - bearerAuth: []
      requestBody:
        required: true
        content:
          application/json:
            schema:
              $ref: "#/components/schemas/AgentPortfolioRequest"
      responses:
        "200":
          description: Portfolio successfully returned
          content:
            application/json:
              schema:
                $ref: "#/components/schemas/AgentPortfolioResponse"
        "403":
          description: Caller cannot access requested agent
        "409":
          description: Refresh is in progress
components:
  schemas:
    AgentPortfolioRequest:
      type: object
      required: [agentCode, timeWindows]
      properties:
        agentCode: { type: string, example: BC00012345 }
        asOfDate: { type: string, format: date, example: "2026-05-19" }
        financialYear: { type: string, example: "2026-27" }
        financialYearStartMonth: { type: integer, default: 4 }
        filters:
          $ref: "#/components/schemas/AgentPortfolioFilters"
        includeNarratives: { type: boolean, default: true }
        includeBenchmarks: { type: boolean, default: true }
        includeFrontendHints: { type: boolean, default: true }
        timeWindows:
          type: array
          items:
            enum:
              - LAST_7_DAYS
              - CURRENT_MONTH
              - PREVIOUS_MONTH_1
              - PREVIOUS_MONTH_2
              - CURRENT_QUARTER
              - CURRENT_FY
              - FY_QUARTER_COMPARISON
    AgentPortfolioFilters:
      type: object
      properties:
        channels:
          type: array
          items: { type: string }
          example: [ONUS, OFFUS]
        transactionTypes:
          type: array
          items: { type: string }
          example: [AEPS, RUPAY, SHG, TPD, IMPS]
        serviceTypes:
          type: array
          items: { type: string }
          example: [CD, CW, FT]
        enrollmentTypes:
          type: array
          items: { type: string }
          example: [PMJJBY, PMSBY, APY]
        leadServiceTypes:
          type: array
          items: { type: string }
          example: [Loan, RD, FD]
        easeServiceTypes:
          type: array
          items: { type: string }
          example: [Debit Card Hotlisting, Account Enquiry]
    AgentPortfolioResponse:
      type: object
      required:
        - requestId
        - schemaVersion
        - referenceDataVersion
        - generatedAt
        - agent
        - portfolioStatus
        - timeWindows
      properties:
        requestId: { type: string }
        schemaVersion: { type: string, example: "2.0" }
        referenceDataVersion: { type: string }
        generatedAt: { type: string, format: date-time }
        agent: { $ref: "#/components/schemas/Agent" }
        portfolioStatus: { $ref: "#/components/schemas/PortfolioStatus" }
        timeWindows:
          type: array
          items: { $ref: "#/components/schemas/AgentPortfolioWindow" }
        recommendations:
          type: array
          items: { $ref: "#/components/schemas/Recommendation" }
    Agent:
      type: object
      properties:
        agentCode: { type: string }
        displayName: { type: string }
        status: { type: string }
        branchCode: { type: string }
        regionCode: { type: string }
        hierarchyLevel: { type: string }
    PortfolioStatus:
      type: object
      properties:
        overallBand: { type: string }
        dailyRefreshStatus: { type: string }
        dataFreshness:
          type: object
          properties:
            latestTxnDate: { type: string, format: date }
            latestCommissionDate: { type: string, format: date }
            latencyMinutes: { type: integer }
        warning:
          type: object
          nullable: true
    Money:
      type: object
      required: [value, currency]
      properties:
        value: { type: number }
        currency: { type: string, example: INR }
    AgentPortfolioWindow:
      type: object
      properties:
        code: { type: string }
        label: { type: string }
        fromDate: { type: string, format: date }
        toDate: { type: string, format: date }
        grain: { type: string }
        periodStatus: { type: string }
        transactionSummary: { $ref: "#/components/schemas/TransactionSummary" }
        commissionSummary: { $ref: "#/components/schemas/CommissionSummary" }
        systemPerformance: { $ref: "#/components/schemas/SystemPerformance" }
        targetCampaignManagement: { $ref: "#/components/schemas/TargetCampaignManagement" }
        anomalyRiskManagement: { $ref: "#/components/schemas/AnomalyRiskManagement" }
        agentAuditManagement: { $ref: "#/components/schemas/AgentAuditManagement" }
        analyticsReportingManagement: { $ref: "#/components/schemas/AnalyticsReportingManagement" }
        enterpriseAgentPerformanceRiskManagement:
          $ref: "#/components/schemas/EAPRMScore"
    TransactionSummary:
      type: object
      properties:
        financialTxnCount: { type: integer }
        nonFinancialTxnCount: { type: integer }
        enrollmentCount: { type: integer }
        leadServiceCount: { type: integer }
        easeServiceCount: { type: integer }
        totalTxnCount: { type: integer }
        amountInvolved: { $ref: "#/components/schemas/Money" }
        financialTxnRatio: { type: number }
        averageFinancialTxnRatio: { type: number }
        cashInflow: { $ref: "#/components/schemas/Money" }
        cashOutflow: { $ref: "#/components/schemas/Money" }
        breakdowns: { type: object }
        series:
          type: array
          items: { type: object }
        narrative: { type: string }
    CommissionSummary:
      type: object
      properties:
        totalCommissionEarned: { $ref: "#/components/schemas/Money" }
        totalEligibleTxnCount: { type: integer }
        eligibleTxnAmount: { $ref: "#/components/schemas/Money" }
        commissionRank: { type: object }
        totalTxnCount: { type: integer }
        averageCommissionEarned: { $ref: "#/components/schemas/Money" }
        narrative: { type: string }
    SystemPerformance:
      type: object
      properties:
        failureRatio: { type: number }
        successRatio: { type: number }
        businessDeclineCount: { type: integer }
        technicalDeclineCount: { type: integer }
        topBusinessDeclineResponseCode: { type: object }
        topTechnicalDeclineResponseCode: { type: object }
        band: { type: string }
        narrative: { type: string }
    TargetCampaignManagement:
      type: object
      properties:
        totalTargetsAssigned: { type: integer }
        totalTargetAchievement: { type: number }
        targets:
          type: array
          items: { type: object }
        narrative: { type: string }
    AnomalyRiskManagement:
      type: object
      properties:
        suspiciousTxnRatio: { type: number }
        totalSuspiciousTxnCount: { type: integer }
        totalAnomalyCaseCount: { type: integer }
        suspiciousCustomerCount: { type: integer }
        suspiciousAmount: { $ref: "#/components/schemas/Money" }
        topAnomalyCase: { type: object }
        riskBand: { type: string }
        narrative: { type: string }
    AgentAuditManagement:
      type: object
      properties:
        totalAuditsInitiated: { type: integer }
        auditedMarksScored: { type: number }
        ongoingAuditCount: { type: integer }
        lastAuditedOn: { type: string, format: date }
        recentAudits:
          type: array
          items: { type: object }
        narrative: { type: string }
    AnalyticsReportingManagement:
      type: object
      properties:
        recommendations:
          type: array
          items: { $ref: "#/components/schemas/Recommendation" }
    EAPRMScore:
      type: object
      properties:
        iScore: { type: integer }
        band: { type: string }
        operationalEfficiencyScore: { type: number }
        behaviouralEfficiencyScore: { type: number }
        complianceScore: { type: number }
        financialHealthScore: { type: number }
        totalCustomersEngaged: { type: integer }
        attendancePercentage: { type: number }
        narrative: { type: string }
    Recommendation:
      type: object
      properties:
        module: { type: string }
        type: { type: string }
        priority: { type: string }
        severity: { type: string }
        message: { type: string }"""

GRAPHQL_SCHEMA = """scalar Date
scalar DateTime
scalar Decimal

type Query {
  agentPortfolio(input: AgentPortfolioInput!): AgentPortfolio!
}

input AgentPortfolioInput {
  agentCode: ID!
  asOfDate: Date
  financialYear: String
  financialYearStartMonth: Int = 4
  timeWindows: [TimeWindowCode!]!
  filters: AgentPortfolioFilterInput
  includeNarratives: Boolean = true
  includeBenchmarks: Boolean = true
  includeFrontendHints: Boolean = true
}

input AgentPortfolioFilterInput {
  channels: [String!]
  transactionTypes: [String!]
  serviceTypes: [String!]
  enrollmentTypes: [String!]
  leadServiceTypes: [String!]
  easeServiceTypes: [String!]
}

enum TimeWindowCode {
  LAST_7_DAYS
  CURRENT_MONTH
  PREVIOUS_MONTH_1
  PREVIOUS_MONTH_2
  CURRENT_QUARTER
  CURRENT_FY
  FY_QUARTER_COMPARISON
}

type AgentPortfolio {
  requestId: ID!
  schemaVersion: String!
  referenceDataVersion: String!
  generatedAt: DateTime!
  agent: Agent!
  portfolioStatus: PortfolioStatus!
  timeWindows: [AgentPortfolioWindow!]!
  recommendations: [Recommendation!]!
}

type Agent {
  agentCode: ID!
  displayName: String!
  status: String!
  branchCode: String
  regionCode: String
  hierarchyLevel: String
}

type PortfolioStatus {
  overallBand: String!
  dailyRefreshStatus: String!
  dataFreshness: DataFreshness
  warning: PortfolioWarning
}

type DataFreshness {
  latestTxnDate: Date
  latestCommissionDate: Date
  latencyMinutes: Int
}

type PortfolioWarning {
  code: String!
  message: String!
  affectedModules: [String!]!
}

type AgentPortfolioWindow {
  code: TimeWindowCode!
  label: String!
  fromDate: Date!
  toDate: Date!
  grain: String!
  periodStatus: String!
  transactionSummary: TransactionSummary!
  commissionSummary: CommissionSummary!
  systemPerformance: SystemPerformance!
  targetCampaignManagement: TargetCampaignManagement!
  anomalyRiskManagement: AnomalyRiskManagement!
  agentAuditManagement: AgentAuditManagement!
  analyticsReportingManagement: AnalyticsReportingManagement!
  enterpriseAgentPerformanceRiskManagement: EAPRMScore!
}

type Money {
  value: Decimal!
  currency: String!
}

type TransactionSummary {
  financialTxnCount: Int!
  nonFinancialTxnCount: Int!
  enrollmentCount: Int!
  leadServiceCount: Int!
  easeServiceCount: Int!
  totalTxnCount: Int!
  amountInvolved: Money!
  financialTxnRatio: Decimal!
  averageFinancialTxnRatio: Decimal
  cashInflow: Money!
  cashOutflow: Money!
  breakdowns: TransactionBreakdowns
  series: [TransactionSeriesPoint!]!
  narrative: String
}

type TransactionBreakdowns {
  byChannel: [BreakdownMetric!]!
  byTransactionType: [BreakdownMetric!]!
  byServiceType: [BreakdownMetric!]!
  byEnrollmentType: [BreakdownMetric!]!
  byLeadServiceType: [BreakdownMetric!]!
  byEaseServiceType: [BreakdownMetric!]!
}

type BreakdownMetric {
  code: String!
  label: String
  totalTxnCount: Int
  amountInvolved: Decimal
  ratio: Decimal
}

type TransactionSeriesPoint {
  date: Date
  month: String
  quarter: String
  financialTxnCount: Int
  nonFinancialTxnCount: Int
  totalTxnCount: Int
  amountInvolved: Decimal
  successRatio: Decimal
  suspiciousTxnRatio: Decimal
}

type CommissionSummary {
  totalCommissionEarned: Money!
  totalEligibleTxnCount: Int!
  eligibleTxnAmount: Money!
  commissionRank: CommissionRank
  totalTxnCount: Int!
  averageCommissionEarned: Money
  narrative: String
}

type CommissionRank {
  rank: Int!
  peerGroup: String!
  movement: Int
}

type SystemPerformance {
  failureRatio: Decimal!
  successRatio: Decimal!
  businessDeclineCount: Int!
  technicalDeclineCount: Int!
  topBusinessDeclineResponseCode: ResponseCodeMetric
  topTechnicalDeclineResponseCode: ResponseCodeMetric
  band: String
  narrative: String
}

type ResponseCodeMetric {
  code: String!
  label: String
  count: Int!
}

type TargetCampaignManagement {
  totalTargetsAssigned: Int!
  totalTargetAchievement: Decimal!
  targets: [TargetMetric!]!
  narrative: String
}

type TargetMetric {
  targetId: ID!
  targetName: String!
  achievementPercent: Decimal!
  targetStatus: String!
  txnPerDayRequired: Int
  pendingTxnCount: Int
  remainingDays: Int
}

type AnomalyRiskManagement {
  suspiciousTxnRatio: Decimal!
  totalSuspiciousTxnCount: Int!
  totalAnomalyCaseCount: Int!
  suspiciousCustomerCount: Int!
  suspiciousAmount: Money!
  topAnomalyCase: AnomalyCaseMetric
  riskBand: String
  narrative: String
}

type AnomalyCaseMetric {
  code: String!
  label: String
  count: Int!
  severity: String
}

type AgentAuditManagement {
  totalAuditsInitiated: Int!
  auditedMarksScored: Decimal
  ongoingAuditCount: Int!
  lastAuditedOn: Date
  recentAudits: [RecentAudit!]!
  narrative: String
}

type RecentAudit {
  auditId: ID!
  auditName: String!
  auditSource: String!
  auditType: String!
  marksScored: Decimal
  status: String!
}

type AnalyticsReportingManagement {
  recommendations: [Recommendation!]!
}

type Recommendation {
  module: String!
  type: String!
  priority: String
  severity: String
  message: String!
}

type EAPRMScore {
  iScore: Int!
  band: String!
  operationalEfficiencyScore: Decimal!
  behaviouralEfficiencyScore: Decimal!
  complianceScore: Decimal!
  financialHealthScore: Decimal!
  totalCustomersEngaged: Int!
  attendancePercentage: Decimal!
  narrative: String
}"""

GRAPHQL_QUERY = """query AgentPortfolio($input: AgentPortfolioInput!) {
  agentPortfolio(input: $input) {
    requestId
    generatedAt
    agent { agentCode displayName status branchCode regionCode }
    portfolioStatus { overallBand dailyRefreshStatus }
    timeWindows {
      code
      fromDate
      toDate
      transactionSummary {
        totalTxnCount
        financialTxnRatio
        amountInvolved { value currency }
        narrative
      }
      commissionSummary {
        totalCommissionEarned { value currency }
        commissionRank { rank peerGroup movement }
        narrative
      }
      enterpriseAgentPerformanceRiskManagement {
        iScore
        band
        operationalEfficiencyScore
        complianceScore
        attendancePercentage
        narrative
      }
    }
    recommendations { module type priority message }
  }
}"""

INTEGRATION_FLOW = """+-------------------------------+       +-------------------------------+
| Agent Portfolio Frontend      |       | Admin / Supervisor Dashboard  |
| Web, mobile, assisted channel |       | Branch, region, enterprise    |
+---------------+---------------+       +---------------+---------------+
                |                                       |
                | REST / GraphQL                        | REST / GraphQL
                v                                       v
        +-------+---------------------------------------+-------+
        | API Gateway / Backend for Frontend (BFF)              |
        | Auth, hierarchy scope, cache, schema versioning        |
        +-------+-------------------+-------------------+-------+
                |                   |                   |
                v                   v                   v
  +-------------+---+     +---------+-------+   +-------+--------------+
  | Agent Master    |     | Portfolio       |   | Reference Data       |
  | Hierarchy       |     | Aggregation     |   | Channel, txn, svc    |
  +-----------------+     +---------+-------+   +----------------------+
                                    |
        +---------------------------+------------------------------+
        |                           |                              |
        v                           v                              v
 +------+-------+        +----------+----------+        +----------+-------+
 | SBOS Autopay |        | iBPM Modules        |        | EAPRM AI/ML      |
 | Commission   |        | TCM ARM SPM AAM ADRM|        | iScore Engine    |
 +--------------+        +---------------------+        +------------------+"""

FRONTEND_WIREFRAME = """+--------------------------------------------------------------------------------+
| Agent: BC00012345  Branch: BR001  Status: ACTIVE       Refresh: Completed       |
+--------------------------------------------------------------------------------+
| iScore 842 Stable | Commission INR 12.4K | Success 97.1% | Suspicious 1.65%     |
+--------------------------------------------------------------------------------+
| Last 7 Days | Current Month | Prev Month | Quarter | Current FY | FY Quarters   |
+--------------------------------------------------------------------------------+
| Transaction Trend Chart              | Target Progress and Daily Required Pace  |
| Cash In / Cash Out / Channel Split   | Pending Txn Count / Remaining Days      |
+--------------------------------------------------------------------------------+
| Risk and Anomaly Panel               | System Performance Panel                 |
| Top anomaly, suspicious customers    | BD/TD codes, failure ratio, narrative    |
+--------------------------------------------------------------------------------+
| Audit Timeline                       | ADRM Recommendations                      |
| Last audit, open audits, marks       | Descriptive / Diagnostic / Predictive    |
+--------------------------------------------------------------------------------+"""


def set_margins(section, top=0.55, bottom=0.55, left=0.6, right=0.6) -> None:
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section, top=0.5, bottom=0.5, left=0.45, right=0.45)


def set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    if section.page_width > section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 8.8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple | list],
    widths: list[float] | None = None,
    font_size: float = 8.8,
    header_fill: str = NAVY,
    zebra: bool = True,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, color=WHITE, size=font_size)
        shade_cell(cell, header_fill)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            set_cell_text(cells[col_idx], str(value), size=font_size)
            if zebra and row_idx % 2 == 1:
                shade_cell(cells[col_idx], "F8FAFC")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_note_box(doc: Document, title: str, body: str, fill: str = LIGHT_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    r.font.size = Pt(10)
    p.add_run("\n" + body)
    doc.add_paragraph()


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def wrap_code_lines(code: str, width: int = 92) -> list[str]:
    wrapped: list[str] = []
    for raw_line in code.strip("\n").splitlines():
        line = raw_line.rstrip()
        if len(line) <= width:
            wrapped.append(line)
            continue
        indent = line[: len(line) - len(line.lstrip(" "))]
        subsequent = indent + "  "
        parts = textwrap.wrap(
            line,
            width=width,
            initial_indent="",
            subsequent_indent=subsequent,
            break_long_words=False,
            break_on_hyphens=False,
        )
        wrapped.extend(parts or [""])
    return wrapped


def add_code(doc: Document, code: str, title: str | None = None) -> None:
    if title:
        add_para(doc, title)
    for line in wrap_code_lines(code):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.86
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        shade_paragraph(p, LIGHT_GRAY)
        run = p.add_run(line or " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(6.6)
        run.font.color.rgb = RGBColor.from_string("111827")
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.8)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(NAVY if style_name != "Heading 3" else TEAL)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Blueprnt Agent Portfolio API Specification | Version 2.0"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
        footer = section.footer.paragraphs[0]
        footer.text = "Confidential solution draft | Prepared 19 May 2026"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)


def build_markdown() -> str:
    lines: list[str] = [
        "# Blueprnt Agent Portfolio API Specification",
        "",
        "Version: 2.0",
        "Prepared on: 19 May 2026",
        "Source reference: Blueprnt API (2).docx",
        "",
        "## Executive Summary",
        "This document defines one REST API and one GraphQL API for complete agent portfolio retrieval by agent code. The API returns module-wise KPI sections, historical windows, financial-year views, classification bands, and plain-language narratives for agent-facing applications.",
        "",
        "## Modules",
    ]
    for code, name, family, responsibility, widget in MODULES:
        lines.append(f"- **{code} - {name}** ({family}): {responsibility} Frontend: {widget}")
    lines.extend(["", "## Time Windows"])
    for code, label, grain, start, end, rule in TIME_WINDOWS:
        lines.append(f"- `{code}`: {label}; grain `{grain}`; example `{start}` to `{end}`. {rule}")
    lines.extend(["", "## REST Endpoint", "`POST /api/v1/agent-portfolios/query`", "", "## GraphQL Query", "`agentPortfolio(input: AgentPortfolioInput!): AgentPortfolio!`", "", "## Metric Catalog"])
    for module, label, field, aggregation, unit, definition in METRIC_CATALOG:
        lines.append(f"- {module}: `{field}` - {label}; aggregation: {aggregation}; unit: {unit}; {definition}")
    lines.extend(["", "## Classification Matrix"])
    for domain, range_, band, action in CLASSIFICATIONS:
        lines.append(f"- {domain}: {range_} => {band}. {action}")
    lines.extend(["", "## REST Request Example", "```json", REST_REQUEST, "```", "", "## REST Response Example", "```json", REST_RESPONSE, "```"])
    return "\n".join(lines)


def build_docx() -> None:
    doc = Document()
    set_margins(doc.sections[0])
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Blueprnt Agent Portfolio\nAPI Specification")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Single REST and GraphQL contract for daily-refreshable, module-wise, historical agent portfolio intelligence")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    add_table(doc, ["Document Field", "Value"], DOCUMENT_META, widths=[2.1, 5.3], font_size=9.2, header_fill=TEAL)
    add_note_box(
        doc,
        "Design intent",
        "The API is intentionally a single portfolio contract. The frontend should not stitch SBOS Autopay, iBPM, and EAPRM data by calling separate module APIs. A backend aggregation layer must return one explainable portfolio object for the agent.",
        fill=LIGHT_BLUE,
    )

    add_page_break(doc)
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This document specifies a production-grade Agent Portfolio API for Blueprnt. The API accepts an agent code and returns the complete daily-refreshable portfolio of the agent across transaction summaries, commission summaries, system performance, targets, anomaly and risk, audits, analytics recommendations, and EAPRM iScore."
    )
    add_para(
        doc,
        "The portfolio is designed for agent-facing screens at the lowest hierarchy level. Therefore, the response does not stop at raw KPIs. Every major module can return a short narrative that explains what the number means, why it matters, and what action the agent should take next."
    )
    add_bullets(doc, [
        "One REST endpoint for complete portfolio retrieval.",
        "One GraphQL query for frontend widgets that need partial field selection.",
        "Time-window arrays for last 7 days, current month, two preceding months, current quarter, current FY, and FY quarter comparison.",
        "Module-wise bifurcation for AUTOPAY, TCM, ARM, SPM, AAM, ADRM, and EAPRM.",
        "Extensible filters for channels, transaction types, service types, enrollment types, lead services, and ease services.",
    ])

    add_heading(doc, "2. Purpose and Outcomes", 1)
    add_table(doc, ["Objective", "Outcome"], [
        ("Consolidate agent intelligence", "A single call returns the complete agent portfolio across operational, commission, risk, audit, analytics, and score modules."),
        ("Support daily refresh", "Frontend screens can show current KPI values and data freshness without triggering expensive recomputation."),
        ("Make KPIs understandable", "Narratives translate mathematical KPIs into simple explanations for agent users."),
        ("Enable historical comparison", "Each KPI can be shown across rolling, monthly, quarterly, and financial-year windows."),
        ("Keep taxonomy future-proof", "Transaction and service types are reference-data driven, so future additions/deletions do not require frontend schema redesign."),
    ], widths=[2.5, 4.9], font_size=9.0)

    add_heading(doc, "3. Content Index", 1)
    add_table(doc, ["Section", "Topic"], [
        ("1-3", "Executive summary, purpose, and document index"),
        ("4", "Business glossary"),
        ("5", "Business and module architecture"),
        ("6", "Frontend integration visuals"),
        ("7", "Financial-year and historical time range model"),
        ("8", "Single REST API specification"),
        ("9", "Single GraphQL API specification"),
        ("10", "Canonical response model and data dictionary"),
        ("11", "Module-wise KPI catalog"),
        ("12", "Classification and banding matrix"),
        ("13", "Aggregation and narrative rules"),
        ("14", "Error, security, refresh, and observability standards"),
        ("15", "Frontend implementation guidance"),
        ("Appendix", "OpenAPI contract, JSON examples, GraphQL examples, versioning notes"),
    ], widths=[1.5, 5.9], font_size=9.0)

    add_heading(doc, "4. Business Glossary", 1)
    add_table(doc, ["Term", "Definition"], GLOSSARY, widths=[1.8, 5.6], font_size=8.8)

    add_page_break(doc)
    add_heading(doc, "5. Business and Module Architecture", 1)
    add_para(
        doc,
        "The API response is organized by business module, not by physical source system. Source systems may be warehouses, transactional stores, model outputs, commission engines, or workflow systems, but the client receives a stable module-wise portfolio contract."
    )
    add_table(doc, ["Code", "Module", "Family", "Business Responsibility", "Frontend Representation"], MODULES, widths=[0.75, 1.7, 1.5, 2.55, 2.45], font_size=7.7, header_fill=NAVY)

    add_heading(doc, "5.1 API Design Principles", 2)
    add_table(doc, ["Principle", "Specification Decision"], [
        ("Single portfolio contract", "One request returns all portfolio modules for an agent and requested time windows."),
        ("Same business logic for REST and GraphQL", "REST and GraphQL must call the same aggregation service to avoid data drift."),
        ("Reference-data driven taxonomy", "Channels, transaction types, and service types are configurable master data, not hard-coded client enums."),
        ("Explainable by default", "Every KPI group can include a narrative and band so agents understand the meaning of the numbers."),
        ("Historical arrays, not flat totals only", "Each module supports period arrays and series data for trend charts and FY comparison."),
        ("Graceful partial data", "If a module source is delayed, return dataQuality warnings and last successful snapshot where permitted."),
    ], widths=[2.1, 5.3], font_size=8.8)

    add_heading(doc, "6. Visual Representation After Frontend Integration", 1)
    add_para(doc, "Figure 1 shows the recommended integration flow from frontend applications to portfolio aggregation and source modules.")
    add_code(doc, INTEGRATION_FLOW)
    add_para(doc, "Figure 2 shows a frontend layout that maps directly to this API response structure.")
    add_code(doc, FRONTEND_WIREFRAME)

    add_heading(doc, "7. Financial-Year and Historical Time Range Model", 1)
    add_note_box(
        doc,
        "Example basis",
        "This document uses asOfDate = 2026-05-19 and India financial year FY 2026-27, running from 2026-04-01 to 2027-03-31. If a tenant uses a different FY start month, the same codes remain valid but dates are recalculated.",
        fill=LIGHT_TEAL,
    )
    add_table(doc, ["Code", "Label", "Grain", "Example From", "Example To", "Rule"], TIME_WINDOWS, widths=[1.55, 1.55, 1.1, 1.05, 1.05, 3.0], font_size=7.7)
    add_bullets(doc, [
        "LAST_7_DAYS must return exactly seven daily records ending on asOfDate.",
        "CURRENT_MONTH and CURRENT_QUARTER are to-date windows, not full future periods.",
        "PREVIOUS_MONTH_1 and PREVIOUS_MONTH_2 are complete historical months.",
        "CURRENT_FY is year-to-date from FY start to asOfDate.",
        "FY_QUARTER_COMPARISON returns Q1-Q4 buckets for the selected FY, with periodStatus COMPLETED, ACTIVE, NOT_STARTED, or PARTIAL.",
    ])

    add_page_break(doc)
    add_heading(doc, "8. Single REST API Specification", 1)
    add_table(doc, ["Property", "Value"], [
        ("Endpoint", "POST /api/v1/agent-portfolios/query"),
        ("Protocol", "HTTPS"),
        ("Content-Type", "application/json"),
        ("Authentication", "OAuth2 bearer token, JWT, or signed service token as per enterprise gateway standard"),
        ("Authorization", "Caller must be mapped to the requested agent through hierarchy, tenant, region, branch, or assigned scope."),
        ("Idempotency", "Read-only. Same request returns same snapshot until source refresh or reference-data version changes."),
        ("Recommended cache key", "agentCode + asOfDate + financialYear + timeWindows + filters + referenceDataVersion"),
        ("SLA target", "p95 <= 800 ms for cached snapshot; p95 <= 2500 ms for computed snapshot."),
    ], widths=[1.9, 5.5], font_size=8.8, header_fill=TEAL)

    add_heading(doc, "8.1 REST Request Data Dictionary", 2)
    add_table(doc, ["Field", "Type", "Required", "Example", "Description"], REQUEST_FIELDS, widths=[1.55, 1.15, 0.8, 1.7, 3.1], font_size=7.8)

    add_heading(doc, "8.2 REST Request Example", 2)
    add_code(doc, REST_REQUEST)
    add_heading(doc, "8.3 REST Root Response Data Dictionary", 2)
    add_table(doc, ["Field", "Type", "Required", "Description"], REST_RESPONSE_FIELD_MAP, widths=[2.4, 1.2, 0.75, 3.0], font_size=7.8)

    complete_map = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(complete_map)
    add_header_footer(doc)
    add_heading(doc, "8.4 Complete REST Module Field Map", 2)
    add_para(doc, "This table lists the complete module-wise response surface. The JSON example that follows is a populated example of the same contract.")
    add_table(doc, ["Module / Object", "Field", "Type", "Meaning"], MODULE_RESPONSE_FIELD_MAP, widths=[1.8, 2.6, 1.25, 5.0], font_size=7.2, header_fill=TEAL)

    rest_example_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(rest_example_section)
    add_header_footer(doc)
    add_heading(doc, "8.5 REST Response Example", 2)
    add_note_box(
        doc,
        "Readable example formatting",
        "The example below is hard-wrapped for Word/Pages readability. The same exact API text is also exported as separate JSON, YAML, and GraphQL files in the docs folder.",
        fill=LIGHT_BLUE,
    )
    add_code(doc, REST_RESPONSE)

    add_page_break(doc)
    add_heading(doc, "9. Single GraphQL API Specification", 1)
    add_para(
        doc,
        "GraphQL is recommended for frontend widgets that do not need the full REST payload. The GraphQL resolver must call the same portfolio aggregation service as REST and must apply the same authorization, filtering, refresh, and classification rules."
    )
    add_heading(doc, "9.1 Complete GraphQL Schema", 2)
    add_code(doc, GRAPHQL_SCHEMA)
    add_heading(doc, "9.2 GraphQL Query Example", 2)
    add_code(doc, GRAPHQL_QUERY)

    add_heading(doc, "10. Canonical Response Model", 1)
    add_table(doc, ["Root Field", "Type", "Description"], ROOT_RESPONSE_FIELDS, widths=[1.8, 1.4, 4.2], font_size=8.8)
    add_heading(doc, "10.1 AgentPortfolioWindow Object", 2)
    add_table(doc, ["Field", "Description"], [
        ("code, label, fromDate, toDate, grain", "Time-window identity and display metadata."),
        ("periodStatus", "COMPLETED, ACTIVE, PARTIAL, NOT_STARTED, or UNAVAILABLE."),
        ("transactionSummary", "Transaction, amount, ratio, channel, service, and series metrics."),
        ("commissionSummary", "SBOS Autopay commission and eligibility metrics."),
        ("systemPerformance", "SPM success, failure, BD, TD, and response-code metrics."),
        ("targetCampaignManagement", "TCM assigned target, achievement, pending count, required daily pace, and target status."),
        ("anomalyRiskManagement", "ARM suspicious transaction, customer, amount, anomaly, and risk band metrics."),
        ("agentAuditManagement", "AAM audit lifecycle, marks, open audits, and recent audit list."),
        ("analyticsReportingManagement", "ADRM descriptive, diagnostic, predictive, and prescriptive recommendations."),
        ("enterpriseAgentPerformanceRiskManagement", "EAPRM iScore and component score metrics."),
    ], widths=[2.3, 5.1], font_size=8.8)

    metric_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(metric_section)
    add_header_footer(doc)
    add_heading(doc, "11. Module-Wise KPI Catalog", 1)
    add_para(doc, "The following table is the implementation-level metric dictionary. It defines what each module must return, how the number should be aggregated, and what unit should be displayed.")
    add_table(doc, ["Module", "Metric", "REST / GraphQL Field", "Aggregation", "Unit", "Business Definition"], METRIC_CATALOG, widths=[1.0, 1.75, 1.9, 1.45, 1.0, 4.3], font_size=7.2, header_fill=NAVY)

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(portrait)
    add_header_footer(doc)
    add_heading(doc, "12. Classification and Banding Matrix", 1)
    add_table(doc, ["Domain", "Range", "Band", "Frontend / Operational Action"], CLASSIFICATIONS, widths=[1.6, 1.1, 1.4, 3.3], font_size=8.1)

    add_heading(doc, "13. Aggregation and Narrative Rules", 1)
    add_table(doc, ["KPI Type", "Aggregation Rule", "Narrative Rule"], [
        ("Counts", "SUM across selected filters and time bucket.", "Mention movement, dominant category, or whether volume is above/below usual behavior."),
        ("Amounts", "SUM with explicit currency object.", "Explain whether value is increasing, decreasing, or concentrated in a channel/service."),
        ("Ratios", "Calculate from raw numerator/denominator, not average of displayed percentages.", "Explain band and operational meaning."),
        ("Ranks", "Return latest rank and movement from previous comparable period.", "Explain rank improvement or decline in peer group."),
        ("Top code/case", "Mode by count; tie-break by amount, severity, then latest occurrence.", "Explain most common reason in plain language."),
        ("Scores", "Return latest score, band, and optional trend movement.", "Explain whether score indicates risk, stability, or model behavior."),
        ("Recommendations", "Generate from rules, model output, and module context.", "Use action-oriented wording: what to do next and why."),
    ], widths=[1.3, 3.0, 3.1], font_size=8.5)
    add_note_box(
        doc,
        "Narrative quality rule",
        "Narratives must be short, non-technical, and action oriented. Example: 'Increase AEPS cash withdrawal by 10 transactions per day to meet the campaign target' is better than 'Achievement ratio is 68.4 percent.'",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "14. Error, Security, Refresh, and Observability Standards", 1)
    add_heading(doc, "14.1 Error Contract", 2)
    add_table(doc, ["HTTP", "Code", "Meaning", "Client Action"], ERRORS, widths=[0.65, 1.8, 2.65, 2.3], font_size=8.2)

    add_heading(doc, "14.2 Security and Compliance", 2)
    add_bullets(doc, [
        "Authorize every request using tenant, hierarchy, branch, region, and agent-scope mapping.",
        "Do not expose customer personally identifiable information in this portfolio response. Use counts, amounts, and risk summaries unless a separate privileged drill-down API is approved.",
        "Mask or hash agent identifiers in logs when required by enterprise logging standards.",
        "Encrypt in transit with TLS 1.2 or later and encrypt portfolio snapshots at rest.",
        "Maintain RBI, UIDAI, NPCI, and bank policy alignment for Aadhaar, payment, risk, and audit data.",
        "Include requestId in every response and every downstream log entry for traceability.",
    ])

    add_heading(doc, "14.3 Daily Refresh Contract", 2)
    add_table(doc, ["Refresh Field", "Meaning"], [
        ("dailyRefreshStatus", "COMPLETED, RUNNING, FAILED, PARTIAL, or SKIPPED."),
        ("latestTxnDate", "Most recent transaction business date included in portfolio computation."),
        ("latestCommissionDate", "Most recent commission business date included from SBOS Autopay."),
        ("latencyMinutes", "Difference between generatedAt and latest available source event/snapshot."),
        ("warning", "Human-readable warning when stale or partial data is returned."),
    ], widths=[2.0, 5.4], font_size=8.7)

    add_heading(doc, "15. Frontend Implementation Guidance", 1)
    add_table(doc, ["Frontend Area", "API Fields", "Recommended Display"], [
        ("Header", "agent, portfolioStatus, dataFreshness", "Agent identity, status, overall band, refresh timestamp."),
        ("KPI strip", "iScore, totalCommissionEarned, successRatio, suspiciousTxnRatio, totalTxnCount", "Compact KPI cards with band color and short narrative."),
        ("Timeline tabs", "timeWindows.code, label, fromDate, toDate", "Tabs or segmented control for historical views."),
        ("Transaction chart", "transactionSummary.series and breakdowns", "Line/bar chart with channel and service filters."),
        ("Target panel", "targetCampaignManagement.targets", "Progress bar, pending transaction count, required daily pace."),
        ("Risk panel", "anomalyRiskManagement", "Risk band, suspicious amount, top anomaly case, investigation priority."),
        ("Audit timeline", "agentAuditManagement.recentAudits", "Newest-first audit history with marks and source."),
        ("Recommendations", "analyticsReportingManagement.recommendations", "Grouped cards: descriptive, diagnostic, predictive, prescriptive."),
    ], widths=[1.6, 3.0, 2.8], font_size=8.2)

    add_heading(doc, "16. Alphabetical Index", 1)
    add_table(doc, ["Index Term", "Reference"], INDEX_TERMS, widths=[2.0, 5.4], font_size=8.5, header_fill=TEAL)

    add_page_break(doc)
    add_heading(doc, "Appendix A. OpenAPI Contract", 1)
    add_code(doc, OPENAPI_EXCERPT)

    add_heading(doc, "Appendix B. Versioning and Extensibility", 1)
    add_bullets(doc, [
        "Use semantic schemaVersion values such as 2.0, 2.1, and 3.0.",
        "Add new fields as optional fields first. Do not remove or rename fields without a major version.",
        "Use referenceDataVersion for changing transaction, channel, enrollment, lead, and ease service taxonomy.",
        "Clients must ignore unknown fields so the response can evolve safely.",
        "When classification thresholds change, publish ruleVersion in portfolioStatus or module metadata.",
    ])

    add_heading(doc, "Appendix C. Implementation Checklist", 1)
    add_table(doc, ["Area", "Checklist Item"], [
        ("API", "REST endpoint implemented with request validation and stable error contract."),
        ("GraphQL", "GraphQL resolver uses the same portfolio aggregation service as REST."),
        ("Security", "Agent hierarchy scope enforced before portfolio computation."),
        ("Data", "Daily snapshot job creates or refreshes portfolio-ready aggregates."),
        ("Taxonomy", "Channels, transaction types, and service types loaded from reference data."),
        ("Frontend", "Widgets consume timeWindows array and handle missing/partial module data."),
        ("Observability", "requestId propagated to gateway, service, database, and downstream logs."),
        ("Testing", "Contract tests cover all time windows, module sections, empty data, and invalid filters."),
    ], widths=[1.5, 5.9], font_size=8.7)

    add_header_footer(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")
    OUT_REST_REQUEST.write_text(REST_REQUEST + "\n", encoding="utf-8")
    OUT_REST_RESPONSE.write_text(REST_RESPONSE + "\n", encoding="utf-8")
    OUT_OPENAPI.write_text(OPENAPI_EXCERPT + "\n", encoding="utf-8")
    OUT_GRAPHQL.write_text(GRAPHQL_SCHEMA + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    print(OUT_DOCX)
    print(OUT_MD)
